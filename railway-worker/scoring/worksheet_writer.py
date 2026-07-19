"""Populate an uploaded HRSA reviewer worksheet while preserving its formatting."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def _insert_after(paragraph: Paragraph, text: str) -> None:
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    created = Paragraph(node, paragraph._parent)
    run = created.add_run(text or "Not identified in the application.")
    run.font.name = "Arial"


def _findings(items: list[dict[str, Any]], finding_type: str = "strength") -> str:
    if not items:
        return "None identified."
    lines = []
    for item in items:
        comment = item.get("comment", "")
        app_pages = item.get("application_pages", item.get("pages", []))
        page_ref = f"(Application p. {', '.join(str(p) for p in app_pages)})" if app_pages else ""
        if finding_type == "weaknesses" and item.get("nofo_requirement"):
            nofo_pages = item.get("nofo_pages", [])
            nofo_ref = f"(NOFO p. {', '.join(str(p) for p in nofo_pages)})" if nofo_pages else ""
            lines.append(f"• {comment} {page_ref}")
            lines.append(f"  NOFO requirement: {item['nofo_requirement']} {nofo_ref}")
            if item.get("impact"):
                lines.append(f"  Impact: {item['impact']}")
        else:
            lines.append(f"• {comment} {page_ref}")
    return "\n".join(lines)


def _normalized(value: str) -> str:
    words = re.sub(r"[^a-z0-9]+", " ", value.lower()).split()
    return " ".join(word for word in words if word not in {"and", "the"})


def populate_reviewer_worksheet(template: Path, output: Path, review: dict[str, Any]) -> Path:
    document = Document(template)
    markers = {"Applicant information:": "applicant_information", "Target population, service area, appropriateness of budget, etc.:": "target_population", "Proposed project/program description:": "project_description", "Major goals and objectives:": "goals_objectives", "Any significant strength and/or weakness:": "significant_findings", "Any other pertinent information:": "other_information"}
    for paragraph in list(document.paragraphs):
        marker = " ".join(paragraph.text.split())
        if marker in markers:
            _insert_after(paragraph, review["overview"].get(markers[marker], ""))
    current = None
    criteria = review["criteria"]
    for paragraph in list(document.paragraphs):
        text = " ".join(paragraph.text.split())
        match = re.match(r"Criterion\s+\d+:\s*(.+?)(?:\s*\(\d+ points\))?$", text, re.I)
        if match:
            heading = match.group(1).strip().lower()
            current = next((c for c in criteria if c["name"].lower() in heading or heading in c["name"].lower()), None)
        if current and text.startswith("Strengths (Please enter"):
            _insert_after(paragraph, _findings(current.get("strengths", []), "strengths"))
        elif current and text.startswith("Mets (Please enter"):
            _insert_after(paragraph, _findings(current.get("mets", []), "mets"))
        elif current and text.startswith("Weaknesses (Please enter"):
            _insert_after(paragraph, _findings(current.get("weaknesses", []), "weaknesses"))
    if len(document.tables) >= 7:
        identity = document.tables[3]
        identity.rows[7].cells[-1].text = review.get("application_number", "")
        identity.rows[8].cells[-1].text = review.get("applicant_name", "")
        score_table = document.tables[4]
        lookup = {_normalized(c["name"]): c for c in criteria}
        sublookup = {_normalized(s["name"]): s for c in criteria for s in c.get("subcriteria", [])}
        for row in score_table.rows[1:-1]:
            label = _normalized(row.cells[1].text)
            item = next((v for k, v in lookup.items() if k in label or label in k), None)
            sub = next((v for k, v in sublookup.items() if k in label or label in k), None)
            if sub: row.cells[-1].text = str(sub["score"])
            elif item: row.cells[-1].text = str(item["score"])
        score_table.rows[-1].cells[-1].text = str(review["final_score"])
        budget = review.get("budget", {})
        if budget.get("recommendation") == "as_requested": document.tables[5].cell(0, 2).text = "X"
        if budget.get("recommendation") == "as_reduced": document.tables[5].cell(0, 4).text = "X"
        years = budget.get("annual_recommended_funding", [])
        for index, amount in enumerate(years[:5], 1):
            if amount is not None: document.tables[6].cell(index, 1).text = f"${amount:,.2f}"
        numeric = [a for a in years[:5] if isinstance(a, (int, float))]
        if numeric: document.tables[6].cell(6, 1).text = f"${sum(numeric):,.2f}"
    for paragraph in list(document.paragraphs):
        if "Rationale for Budget Reduction" in paragraph.text:
            _insert_after(paragraph, review.get("budget", {}).get("reduction_rationale", "Not applicable."))
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return output
