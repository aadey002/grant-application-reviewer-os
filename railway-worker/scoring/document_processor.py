"""
Document Processor for Grant Applications

Handles processing of PDF and Word documents, extracting text, tables, and structure.
Supports grant applications and NOFO documents.
"""

import os
import re
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import logging

# Document processing libraries
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

from datetime import datetime

class DocumentProcessor:
    """Main document processing class for grant applications and NOFOs."""

    def __init__(self):
        self.supported_formats = ['.pdf', '.doc', '.docx']
        self.logger = logging.getLogger(__name__)

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract all relevant information.

        Args:
            file_path: Absolute path to the document file

        Returns:
            Dictionary containing processed document information
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")

        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {extension}")

        if extension == '.pdf':
            return self._process_pdf(file_path)
        elif extension in ['.doc', '.docx']:
            return self._process_word(file_path)
        else:
            raise ValueError(f"Unsupported document type: {extension}")

    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Process PDF document and extract content."""
        result = {
            'file_path': str(file_path),
            'document_type': 'PDF',
            'processed_date': datetime.now().isoformat(),
            'total_pages': 0,
            'text_content': '',
            'sections': {},
            'tables': [],
            'metadata': {},
            'page_contents': []
        }

        try:
            # Use PyMuPDF for text extraction
            if fitz:
                doc = fitz.open(file_path)
                result['total_pages'] = len(doc)
                result['metadata'] = doc.metadata

                full_text = ""
                page_contents = []

                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    full_text += f"\n--- PAGE {page_num + 1} ---\n{page_text}"
                    page_contents.append({
                        'page_number': page_num + 1,
                        'text': page_text,
                        'word_count': len(page_text.split())
                    })

                result['text_content'] = full_text
                result['page_contents'] = page_contents
                doc.close()

            # Use pdfplumber for table extraction
            if pdfplumber:
                with pdfplumber.open(file_path) as pdf:
                    tables = []
                    for page_num, page in enumerate(pdf.pages):
                        page_tables = page.extract_tables()
                        for table_idx, table in enumerate(page_tables):
                            if table:
                                tables.append({
                                    'page': page_num + 1,
                                    'table_index': table_idx,
                                    'data': table,
                                    'rows': len(table),
                                    'columns': len(table[0]) if table else 0
                                })
                    result['tables'] = tables

            # Extract document sections
            result['sections'] = self._extract_sections(result['text_content'])

        except Exception as e:
            self.logger.error(f"Error processing PDF {file_path}: {str(e)}")
            result['error'] = str(e)

        return result

    def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """Process Word document and extract content."""
        result = {
            'file_path': str(file_path),
            'document_type': 'Word',
            'processed_date': datetime.now().isoformat(),
            'total_pages': 0,
            'text_content': '',
            'sections': {},
            'tables': [],
            'metadata': {},
            'paragraphs': []
        }

        try:
            if Document:
                doc = Document(file_path)

                # Extract metadata
                core_props = doc.core_properties
                result['metadata'] = {
                    'title': core_props.title,
                    'author': core_props.author,
                    'created': core_props.created.isoformat() if core_props.created else None,
                    'modified': core_props.modified.isoformat() if core_props.modified else None
                }

                # Extract text content
                full_text = ""
                paragraphs = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text += para.text + "\n"
                        paragraphs.append({
                            'text': para.text,
                            'style': para.style.name if para.style else 'Normal',
                            'word_count': len(para.text.split())
                        })

                result['text_content'] = full_text
                result['paragraphs'] = paragraphs

                # Extract tables
                tables = []
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)

                    tables.append({
                        'table_index': table_idx,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    })

                result['tables'] = tables

                # Estimate page count (rough approximation)
                result['total_pages'] = max(1, len(full_text) // 3000)

                # Extract document sections
                result['sections'] = self._extract_sections(full_text)

        except Exception as e:
            self.logger.error(f"Error processing Word document {file_path}: {str(e)}")
            result['error'] = str(e)

        return result

    def _extract_sections(self, text_content: str) -> Dict[str, Dict[str, Any]]:
        """Extract document sections based on common grant application patterns."""
        sections = {}

        # Common grant application section patterns
        section_patterns = {
            'abstract': r'(?i)(abstract|executive\s+summary|project\s+summary)',
            'statement_of_need': r'(?i)(statement\s+of\s+need|needs?\s+assessment|problem\s+statement)',
            'project_description': r'(?i)(project\s+description|program\s+description|approach)',
            'methodology': r'(?i)(methodology|methods?|approach|implementation)',
            'goals_objectives': r'(?i)(goals?\s+and\s+objectives?|aims?|objectives?)',
            'timeline': r'(?i)(timeline|schedule|work\s+plan|project\s+timeline)',
            'budget': r'(?i)(budget|financial|cost|funding)',
            'evaluation': r'(?i)(evaluation|assessment|outcomes?|monitoring)',
            'sustainability': r'(?i)(sustainability|continuation|long.?term)',
            'organizational_capacity': r'(?i)(organizational\s+capacity|qualifications|experience)',
            'personnel': r'(?i)(personnel|staff|team|biographical|biosketches?)',
            'appendices': r'(?i)(appendix|appendices|attachments?)'
        }

        # Split text into paragraphs
        paragraphs = text_content.split('\n')
        current_section = None
        section_content = []

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if paragraph is a section header
            for section_name, pattern in section_patterns.items():
                if re.match(pattern, para) and len(para) < 100:
                    # Save previous section
                    if current_section and section_content:
                        sections[current_section] = self._process_section_content(section_content)

                    # Start new section
                    current_section = section_name
                    section_content = []
                    break
            else:
                # Add to current section
                if current_section:
                    section_content.append(para)

        # Save last section
        if current_section and section_content:
            sections[current_section] = self._process_section_content(section_content)

        return sections

    def _process_section_content(self, content: List[str]) -> Dict[str, Any]:
        """Process content of a section and extract metadata."""
        full_content = '\n'.join(content)

        return {
            'content': full_content,
            'word_count': len(full_content.split()),
            'paragraph_count': len(content),
            'preview': full_content[:200] + '...' if len(full_content) > 200 else full_content,
            'pages': self._estimate_page_numbers(full_content)
        }

    def _estimate_page_numbers(self, content: str) -> str:
        """Estimate page numbers where content appears (rough approximation)."""
        words = len(content.split())
        start_page = max(1, words // 300)
        end_page = start_page + max(1, words // 250)

        if start_page == end_page:
            return str(start_page)
        else:
            return f"{start_page}-{end_page}"

    def extract_specific_sections(self, document_content: Dict[str, Any],
                                section_names: List[str]) -> Dict[str, Dict[str, Any]]:
        """Extract specific sections from processed document content."""
        all_sections = document_content.get('sections', {})
        requested_sections = {}

        for section_name in section_names:
            # Flexible matching
            for key, content in all_sections.items():
                if section_name.lower() in key.lower():
                    requested_sections[key] = content
                    break

        return requested_sections

    def search_content(self, document_content: Dict[str, Any],
                      search_terms: List[str]) -> List[Dict[str, Any]]:
        """Search for specific terms in document content."""
        results = []
        text_content = document_content.get('text_content', '')

        for term in search_terms:
            pattern = re.compile(re.escape(term), re.IGNORECASE)
            matches = list(pattern.finditer(text_content))

            for match in matches:
                # Find surrounding context
                start = max(0, match.start() - 100)
                end = min(len(text_content), match.end() + 100)
                context = text_content[start:end]

                results.append({
                    'term': term,
                    'position': match.start(),
                    'context': context,
                    'page': self._estimate_page_from_position(match.start(), text_content)
                })

        return results

    def _estimate_page_from_position(self, position: int, full_text: str) -> int:
        """Estimate page number from character position."""
        page_markers = re.findall(r'--- PAGE (\d+) ---', full_text[:position])
        if page_markers:
            return int(page_markers[-1])
        else:
            return max(1, position // 3000)
