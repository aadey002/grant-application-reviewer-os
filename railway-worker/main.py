"""
Grant Application Reviewer — Railway Worker
Production FastAPI backend. All persistence via Supabase (DB + Storage).
No SQLite. No localhost file paths persisted between requests.
"""
from __future__ import annotations

import io
import json
import logging
import os
import re
import sys
import tempfile
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# ---------------------------------------------------------------------------
# Scoring module path — added to sys.path but NOT imported at startup
# ---------------------------------------------------------------------------
sys.path.insert(0, str(Path(__file__).parent / "scoring"))

# Heavy imports (PyMuPDF, anthropic, python-docx) are loaded lazily in job handlers
# to keep startup fast and prevent container crash if a native dep is missing.

from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel

# ---------------------------------------------------------------------------
# Lazy imports for heavy dependencies (PyMuPDF, anthropic, python-docx)
# ---------------------------------------------------------------------------
def _lazy_scoring():
    from scoring.safe_review import extract_nofo_criteria, safe_extract_application_zip
    from scoring.anthropic_review import score_application_with_claude
    from scoring.worksheet_writer import populate_reviewer_worksheet
    return extract_nofo_criteria, safe_extract_application_zip, score_application_with_claude, populate_reviewer_worksheet

# Supabase client — only imported when needed
_supabase_client: Any = None

def get_supabase():
    global _supabase_client
    if _supabase_client is None:
        from supabase import create_client
        url = os.environ.get("SUPABASE_URL", "")
        key = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
        if not url or not key:
            raise RuntimeError("SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY must be set")
        _supabase_client = create_client(url, key)
    return _supabase_client

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s — %(message)s")
logger = logging.getLogger("grant_worker")

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_SERVICE_ROLE_KEY = os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")
ALLOWED_ORIGINS_RAW = os.getenv("ALLOWED_ORIGINS", "")
ALLOWED_ORIGINS: list[str] = [o.strip() for o in ALLOWED_ORIGINS_RAW.split(",") if o.strip()] or ["*"]

# Supabase Storage buckets
BUCKET_NOFO = "nofo-files"
BUCKET_APPS = "grant-applications"
BUCKET_WORKSHEETS = "worksheet-templates"
BUCKET_COMPLETED = "completed-worksheets"

# ---------------------------------------------------------------------------
# App
# ---------------------------------------------------------------------------
app = FastAPI(title="Grant Application Reviewer Worker", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Supabase client defined above as lazy singleton

# ---------------------------------------------------------------------------
# Storage helpers
# ---------------------------------------------------------------------------

def _upload_bytes(sb: Client, bucket: str, path: str, data: bytes, content_type: str = "application/octet-stream") -> str:
    """Upload raw bytes to Supabase Storage; return the storage path."""
    sb.storage.from_(bucket).upload(
        path=path,
        file=data,
        file_options={"content-type": content_type, "upsert": "true"},
    )
    return path


def _download_bytes(sb: Client, bucket: str, path: str) -> bytes:
    """Download a file from Supabase Storage as bytes."""
    return sb.storage.from_(bucket).download(path)


def _signed_url(sb: Client, bucket: str, path: str, expires_in: int = 3600) -> str:
    result = sb.storage.from_(bucket).create_signed_url(path, expires_in)
    return result["signedURL"]


def _delete_storage_prefix(sb: Client, bucket: str, prefix: str) -> None:
    """Delete all objects whose path starts with prefix."""
    try:
        items = sb.storage.from_(bucket).list(prefix)
        paths = [f"{prefix}/{item['name']}" for item in (items or []) if item.get("name")]
        if paths:
            sb.storage.from_(bucket).remove(paths)
    except Exception as exc:
        logger.warning("Storage cleanup error for %s/%s: %s", bucket, prefix, exc)


# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------

def _insert(sb: Client, table: str, row: dict[str, Any]) -> dict[str, Any]:
    result = sb.table(table).insert(row).execute()
    return result.data[0] if result.data else row


def _update(sb: Client, table: str, match: dict[str, Any], values: dict[str, Any]) -> None:
    q = sb.table(table).update(values)
    for col, val in match.items():
        q = q.eq(col, val)
    q.execute()


def _select(sb: Client, table: str, match: dict[str, Any]) -> list[dict[str, Any]]:
    q = sb.table(table).select("*")
    for col, val in match.items():
        q = q.eq(col, val)
    return q.execute().data or []


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.get("/health")
def health():
    return {"status": "ok", "timestamp": datetime.now(timezone.utc).isoformat()}

@app.get("/ready")
def ready():
    checks = {
        "supabase_configured": bool(SUPABASE_URL and SUPABASE_SERVICE_ROLE_KEY),
        "anthropic_configured": bool(os.environ.get("ANTHROPIC_API_KEY")),
        "scoring_mode": "claude",
    }
    ready = all([checks["supabase_configured"], checks["anthropic_configured"]])
    return {"ready": ready, **checks}


# ---------------------------------------------------------------------------
# POST /safe-reviews/extract-rubric
# ---------------------------------------------------------------------------

@app.post("/safe-reviews/extract-rubric")
async def extract_rubric(
    nofo: UploadFile = File(...),
    agency: str = Form("HRSA"),
):
    """Accept a NOFO file upload and return extracted criterion rubric JSON."""
    nofo_bytes = await nofo.read()
    suffix = Path(nofo.filename or "nofo.pdf").suffix.lower() or ".pdf"

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(nofo_bytes)
        tmp_path = Path(tmp.name)

    try:
        extract_nofo_criteria, *_ = _lazy_scoring()
        rubric = extract_nofo_criteria(tmp_path)
    except Exception as exc:
        logger.exception("extract_nofo_criteria failed")
        raise HTTPException(status_code=422, detail=str(exc))
    finally:
        tmp_path.unlink(missing_ok=True)

    return JSONResponse(content={"success": True, "rubric": rubric})


# ---------------------------------------------------------------------------
# POST /nofo-brief/generate
# ---------------------------------------------------------------------------

NOFO_BRIEF_TOOL_SCHEMA = {
    "name": "submit_nofo_brief",
    "description": "Submit a structured NOFO brief extracted from the funding opportunity document.",
    "input_schema": {
        "type": "object",
        "additionalProperties": False,
        "required": [
            "funding_opportunity", "program_purpose", "eligibility", "funding",
            "important_dates", "required_components", "program_requirements",
            "scoring_criteria", "budget_requirements", "special_priorities",
            "compliance_risks", "reviewer_checklist", "executive_summary",
        ],
        "properties": {
            "funding_opportunity": {
                "type": "object", "additionalProperties": False,
                "required": ["title", "number", "agency", "bureau"],
                "properties": {
                    "title": {"type": "string"},
                    "number": {"type": "string"},
                    "agency": {"type": "string"},
                    "bureau": {"type": "string"},
                },
            },
            "program_purpose": {
                "type": "object", "additionalProperties": False,
                "required": ["text", "citations"],
                "properties": {
                    "text": {"type": "string"},
                    "citations": {"type": "array", "items": {
                        "type": "object", "additionalProperties": False,
                        "required": ["page", "text"],
                        "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                    }},
                },
            },
            "eligibility": {
                "type": "object", "additionalProperties": False,
                "required": ["eligible_applicants", "target_populations", "geographic_requirements"],
                "properties": {
                    "eligible_applicants": {
                        "type": "object", "additionalProperties": False,
                        "required": ["text", "citations"],
                        "properties": {
                            "text": {"type": "string"},
                            "citations": {"type": "array", "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["page", "text"],
                                "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                            }},
                        },
                    },
                    "target_populations": {
                        "type": "object", "additionalProperties": False,
                        "required": ["text", "citations"],
                        "properties": {
                            "text": {"type": "string"},
                            "citations": {"type": "array", "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["page", "text"],
                                "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                            }},
                        },
                    },
                    "geographic_requirements": {
                        "type": "object", "additionalProperties": False,
                        "required": ["text", "citations"],
                        "properties": {
                            "text": {"type": "string"},
                            "citations": {"type": "array", "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["page", "text"],
                                "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                            }},
                        },
                    },
                },
            },
            "funding": {
                "type": "object", "additionalProperties": False,
                "required": ["amounts", "award_range", "project_period"],
                "properties": {
                    "amounts": {
                        "type": "object", "additionalProperties": False,
                        "required": ["text", "citations"],
                        "properties": {
                            "text": {"type": "string"},
                            "citations": {"type": "array", "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["page", "text"],
                                "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                            }},
                        },
                    },
                    "award_range": {
                        "type": "object", "additionalProperties": False,
                        "required": ["text", "citations"],
                        "properties": {
                            "text": {"type": "string"},
                            "citations": {"type": "array", "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["page", "text"],
                                "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                            }},
                        },
                    },
                    "project_period": {
                        "type": "object", "additionalProperties": False,
                        "required": ["text", "citations"],
                        "properties": {
                            "text": {"type": "string"},
                            "citations": {"type": "array", "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["page", "text"],
                                "properties": {"page": {"type": "integer", "minimum": 1}, "text": {"type": "string"}},
                            }},
                        },
                    },
                },
            },
            "important_dates": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["date", "description", "nofo_page"],
                    "properties": {
                        "date": {"type": "string"},
                        "description": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "required_components": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["component", "page_limit", "nofo_page"],
                    "properties": {
                        "component": {"type": "string"},
                        "page_limit": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "program_requirements": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["requirement", "nofo_page"],
                    "properties": {
                        "requirement": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "scoring_criteria": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["name", "points", "subcriteria", "nofo_page"],
                    "properties": {
                        "name": {"type": "string"},
                        "points": {"type": "integer", "minimum": 0},
                        "subcriteria": {
                            "type": "array",
                            "items": {
                                "type": "object", "additionalProperties": False,
                                "required": ["name", "points"],
                                "properties": {
                                    "name": {"type": "string"},
                                    "points": {"type": "integer", "minimum": 0},
                                },
                            },
                        },
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "budget_requirements": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["requirement", "nofo_page"],
                    "properties": {
                        "requirement": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "special_priorities": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["priority", "nofo_page"],
                    "properties": {
                        "priority": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "compliance_risks": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["risk", "nofo_page"],
                    "properties": {
                        "risk": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "reviewer_checklist": {
                "type": "array",
                "items": {
                    "type": "object", "additionalProperties": False,
                    "required": ["item", "nofo_page"],
                    "properties": {
                        "item": {"type": "string"},
                        "nofo_page": {"type": "integer", "minimum": 1},
                    },
                },
            },
            "executive_summary": {"type": "string"},
        },
    },
}

NOFO_BRIEF_SYSTEM_PROMPT = """You are an expert federal grant analyst preparing a concise but comprehensive Reviewer NOFO Brief. Your job is to extract and organize all information a reviewer needs to evaluate applications against this NOFO. Cite every statement with the exact NOFO page number (e.g., "NOFO p. 12"). Be precise, exhaustive on requirements, and err on the side of inclusion for compliance risks and reviewer checklist items. Do not paraphrase requirements in ways that change their meaning. Identify all explicit and implicit requirements, important dates, page/budget limits, and scoring criteria exactly as stated."""


def _extract_nofo_text_for_brief(nofo_bytes: bytes) -> tuple[list[str], str]:
    """Extract text from NOFO PDF using PyMuPDF, returning (pages, full_text)."""
    import fitz  # PyMuPDF
    import io as _io
    doc = fitz.open(stream=nofo_bytes, filetype="pdf")
    pages: list[str] = []
    blocks: list[str] = []
    for i, page in enumerate(doc, 1):
        text = page.get_text()
        pages.append(text)
        blocks.append(f"\n--- NOFO PAGE {i} ---\n{text.strip()}")
    doc.close()
    return pages, "".join(blocks)


def _generate_nofo_brief_docx(brief: dict[str, Any]) -> bytes:
    """Generate a DOCX from the structured NOFO brief and return bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    fo = brief.get("funding_opportunity", {})
    title_para = doc.add_heading(fo.get("title", "NOFO Reviewer Brief"), level=0)

    meta_lines = [
        f"Number: {fo.get('number', 'N/A')}",
        f"Agency: {fo.get('agency', 'N/A')}",
        f"Bureau: {fo.get('bureau', 'N/A')}",
    ]
    for line in meta_lines:
        doc.add_paragraph(line)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(brief.get("executive_summary", ""))

    # Program Purpose
    doc.add_heading("Program Purpose", level=1)
    pp = brief.get("program_purpose", {})
    doc.add_paragraph(pp.get("text", ""))
    for c in pp.get("citations", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"NOFO p. {c.get('page', '?')}: ").bold = True
        p.add_run(c.get("text", ""))

    # Eligibility
    doc.add_heading("Eligibility", level=1)
    elig = brief.get("eligibility", {})
    for field_key, field_label in [
        ("eligible_applicants", "Eligible Applicants"),
        ("target_populations", "Target Populations"),
        ("geographic_requirements", "Geographic Requirements"),
    ]:
        doc.add_heading(field_label, level=2)
        field = elig.get(field_key, {})
        doc.add_paragraph(field.get("text", ""))
        for c in field.get("citations", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"NOFO p. {c.get('page', '?')}: ").bold = True
            p.add_run(c.get("text", ""))

    # Funding
    doc.add_heading("Funding", level=1)
    funding = brief.get("funding", {})
    for field_key, field_label in [
        ("amounts", "Total Funding Available"),
        ("award_range", "Award Range"),
        ("project_period", "Project Period"),
    ]:
        doc.add_heading(field_label, level=2)
        field = funding.get(field_key, {})
        doc.add_paragraph(field.get("text", ""))
        for c in field.get("citations", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"NOFO p. {c.get('page', '?')}: ").bold = True
            p.add_run(c.get("text", ""))

    # Important Dates
    doc.add_heading("Important Dates", level=1)
    for item in brief.get("important_dates", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{item.get('date', 'TBD')} — {item.get('description', '')} ").bold = False
        p.add_run(f"(NOFO p. {item.get('nofo_page', '?')})").italic = True

    # Required Components
    doc.add_heading("Required Application Components", level=1)
    for item in brief.get("required_components", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{item.get('component', '')}").bold = True
        limit = item.get("page_limit", "")
        if limit:
            p.add_run(f" — {limit}")
        p.add_run(f" (NOFO p. {item.get('nofo_page', '?')})").italic = True

    # Scoring Criteria
    doc.add_heading("Scoring Criteria", level=1)
    for crit in brief.get("scoring_criteria", []):
        doc.add_heading(f"{crit.get('name', '')} — {crit.get('points', 0)} pts (NOFO p. {crit.get('nofo_page', '?')})", level=2)
        for sub in crit.get("subcriteria", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{sub.get('name', '')}: {sub.get('points', 0)} pts")

    # Program Requirements
    doc.add_heading("Program Requirements", level=1)
    for item in brief.get("program_requirements", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item.get("requirement", ""))
        p.add_run(f" (NOFO p. {item.get('nofo_page', '?')})").italic = True

    # Budget Requirements
    doc.add_heading("Budget Requirements", level=1)
    for item in brief.get("budget_requirements", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item.get("requirement", ""))
        p.add_run(f" (NOFO p. {item.get('nofo_page', '?')})").italic = True

    # Special Priorities
    if brief.get("special_priorities"):
        doc.add_heading("Special Priorities", level=1)
        for item in brief.get("special_priorities", []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item.get("priority", ""))
            p.add_run(f" (NOFO p. {item.get('nofo_page', '?')})").italic = True

    # Compliance Risks
    doc.add_heading("Compliance Risks", level=1)
    for item in brief.get("compliance_risks", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item.get("risk", ""))
        p.add_run(f" (NOFO p. {item.get('nofo_page', '?')})").italic = True

    # Reviewer Checklist
    doc.add_heading("Reviewer Checklist", level=1)
    for item in brief.get("reviewer_checklist", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("[ ] ")
        p.add_run(item.get("item", ""))
        p.add_run(f" (NOFO p. {item.get('nofo_page', '?')})").italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _run_nofo_brief(brief_id: str, review_id: str, nofo_storage_path: str, agency: str, criteria_json: str) -> None:
    """Background task: generate NOFO brief with Claude, store result."""
    sb = get_supabase()
    logger.info("NOFO brief %s starting", brief_id)
    try:
        # Download NOFO from storage
        nofo_bytes = _download_bytes(sb, BUCKET_NOFO, nofo_storage_path)

        # Extract text using PyMuPDF
        pages, nofo_text = _extract_nofo_text_for_brief(nofo_bytes)
        logger.info("NOFO brief %s: extracted %d pages", brief_id, len(pages))

        # Build prompt
        criteria = json.loads(criteria_json) if criteria_json else []
        rubric_hint = ""
        if criteria:
            rubric_hint = "\n\nThe approved scoring rubric is:\n" + "\n".join(
                f"- {c.get('name', '')}: {c.get('points', 0)} pts" for c in criteria
            )

        prompt = (
            f"Agency: {agency}\n\n"
            "You are preparing a Reviewer NOFO Brief. Every statement must be cited with 'NOFO p. X' "
            "where X is the exact page number from the NOFO text below.\n"
            f"{rubric_hint}\n\n"
            "NOFO TEXT:\n"
            f"{nofo_text[:120000]}\n\n"
            "Extract all information needed to complete the brief. Be thorough and precise. "
            "For every field that includes citations, cite the exact NOFO page numbers."
        )

        import anthropic
        client = anthropic.Anthropic(
            api_key=os.getenv("ANTHROPIC_API_KEY", "").strip(),
            timeout=600.0,
        )
        response = client.messages.create(
            model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-5"),
            max_tokens=16000,
            temperature=0,
            system=NOFO_BRIEF_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}],
            tools=[NOFO_BRIEF_TOOL_SCHEMA],
            tool_choice={"type": "tool", "name": "submit_nofo_brief"},
        )

        tool_use = next(
            (b for b in response.content if b.type == "tool_use" and b.name == "submit_nofo_brief"),
            None,
        )
        if not tool_use:
            raise RuntimeError("Claude did not return a structured NOFO brief")

        brief_data = tool_use.input
        if isinstance(brief_data, str):
            brief_data = json.loads(brief_data)

        # Generate DOCX
        docx_bytes = _generate_nofo_brief_docx(brief_data)
        docx_filename = f"nofo_brief_{brief_id}.docx"
        docx_storage_path = f"{review_id}/{docx_filename}"
        _upload_bytes(
            sb, BUCKET_COMPLETED, docx_storage_path, docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        # Update DB row to ready
        _update(sb, "nofo_briefs", {"id": brief_id}, {
            "status": "ready",
            "brief_json": json.dumps(brief_data),
            "docx_storage_path": docx_storage_path,
            "generated_at": datetime.now(timezone.utc).isoformat(),
        })
        logger.info("NOFO brief %s ready", brief_id)

    except Exception as exc:
        logger.exception("NOFO brief %s failed: %s", brief_id, exc)
        _update(sb, "nofo_briefs", {"id": brief_id}, {
            "status": "failed",
            "error_message": str(exc)[:2000],
        })


@app.post("/nofo-brief/generate", status_code=202)
async def generate_nofo_brief(
    review_id: str = Form(...),
    nofo_storage_path: str = Form(...),
    agency: str = Form("HRSA"),
    criteria_json: str = Form(...),
):
    """Generate a Reviewer NOFO Brief from the stored NOFO PDF using Claude."""
    sb = get_supabase()
    brief_id = str(uuid.uuid4())

    _insert(sb, "nofo_briefs", {
        "id": brief_id,
        "review_id": review_id,
        "nofo_storage_path": nofo_storage_path,
        "agency": agency,
        "status": "generating",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    import threading
    def _run_brief_thread():
        try:
            _run_nofo_brief(brief_id, review_id, nofo_storage_path, agency, criteria_json)
        except Exception as exc:
            logger.error("NOFO brief %s failed: %s", brief_id, exc)
            try:
                sb2 = get_supabase()
                _update(sb2, "nofo_briefs", {"id": brief_id}, {"status": "failed", "error_message": str(exc)[:2000]})
            except Exception:
                pass
    threading.Thread(target=_run_brief_thread, daemon=True).start()

    return JSONResponse(status_code=202, content={
        "brief_id": brief_id,
        "review_id": review_id,
        "status": "generating",
        "message": "NOFO brief generation started.",
    })


@app.get("/nofo-brief/{brief_id}")
def get_nofo_brief(brief_id: str):
    """Return the NOFO brief JSON and download URL."""
    sb = get_supabase()
    rows = _select(sb, "nofo_briefs", {"id": brief_id})
    if not rows:
        raise HTTPException(status_code=404, detail="NOFO brief not found")
    row = rows[0]

    result: dict[str, Any] = {
        "brief_id": brief_id,
        "review_id": row.get("review_id"),
        "status": row.get("status"),
        "agency": row.get("agency"),
        "created_at": row.get("created_at"),
        "generated_at": row.get("generated_at"),
        "error_message": row.get("error_message"),
        "brief_json": json.loads(row["brief_json"]) if row.get("brief_json") else None,
        "docx_download_url": None,
    }

    if row.get("docx_storage_path") and row.get("status") == "ready":
        try:
            result["docx_download_url"] = _signed_url(sb, BUCKET_COMPLETED, row["docx_storage_path"], expires_in=3600)
        except Exception as exc:
            logger.warning("Could not generate signed URL for brief %s: %s", brief_id, exc)

    return result


@app.get("/nofo-brief/{brief_id}/download")
def download_nofo_brief(brief_id: str):
    """Return a signed download URL for the NOFO brief DOCX."""
    sb = get_supabase()
    rows = _select(sb, "nofo_briefs", {"id": brief_id})
    if not rows:
        raise HTTPException(status_code=404, detail="NOFO brief not found")
    row = rows[0]
    if row.get("status") != "ready" or not row.get("docx_storage_path"):
        raise HTTPException(status_code=404, detail=f"NOFO brief not ready (status: {row.get('status', 'unknown')})")
    try:
        signed = _signed_url(sb, BUCKET_COMPLETED, row["docx_storage_path"], expires_in=3600)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not generate download URL: {exc}")
    return {
        "brief_id": brief_id,
        "download_url": signed,
        "expires_in_seconds": 3600,
    }


# ---------------------------------------------------------------------------
# POST /safe-reviews/enqueue — lightweight, accepts Supabase Storage paths only
# ---------------------------------------------------------------------------

@app.post("/safe-reviews/enqueue", status_code=202)
async def enqueue_review(
    background_tasks: BackgroundTasks,
    application_storage_paths: str = Form(...),
    nofo_storage_path: str = Form(...),
    rubric_storage_path: Optional[str] = Form(None),
    worksheet_storage_path: Optional[str] = Form(None),
    approved_criteria: str = Form(...),
    agency: str = Form("HRSA"),
    user_id: str = Form(...),
    review_id: str = Form(...),
):
    """Lightweight enqueue: files are already in Supabase Storage. Creates jobs and returns immediately."""
    sb = get_supabase()
    try:
        criteria = json.loads(approved_criteria)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=422, detail=f"Invalid approved_criteria JSON: {exc}")
    try:
        app_paths = json.loads(application_storage_paths)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=422, detail=f"Invalid application_storage_paths JSON: {exc}")

    prefix = f"{user_id}/{review_id}"
    job_records = []
    for idx, app_path in enumerate(app_paths):
        application_id = str(uuid.uuid4())
        job_id = str(uuid.uuid4())
        filename = app_path.rsplit("/", 1)[-1] if "/" in app_path else app_path
        _insert(sb, "applications", {
            "id": application_id, "review_id": review_id, "user_id": user_id,
            "filename": filename, "storage_path": app_path, "status": "queued",
            "agency": agency, "criteria": json.dumps(criteria),
            "nofo_storage_path": nofo_storage_path,
            "worksheet_storage_path": worksheet_storage_path or "",
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        _insert(sb, "processing_jobs", {
            "id": job_id, "application_id": application_id, "review_id": review_id,
            "user_id": user_id, "status": "queued", "agency": agency,
            "criteria": json.dumps(criteria),
            "nofo_storage_path": nofo_storage_path,
            "worksheet_storage_path": worksheet_storage_path or "",
            "created_at": datetime.now(timezone.utc).isoformat(),
        })
        job_records.append({"job_id": job_id, "application_id": application_id,
                           "filename": filename, "status": "queued"})
        background_tasks.add_task(
            _process_job, job_id, app_path, criteria, agency, review_id,
            nofo_storage_path, worksheet_storage_path or "",
        )
    return JSONResponse(status_code=202, content={
        "review_id": review_id, "jobs": job_records,
        "status": "queued", "message": f"{len(job_records)} application(s) queued for processing.",
    })


# ---------------------------------------------------------------------------
# POST /safe-reviews/run — legacy endpoint, accepts file uploads
# ---------------------------------------------------------------------------

@app.post("/safe-reviews/run")
async def run_review(
    background_tasks: BackgroundTasks,
    applications: list[UploadFile] = File(...),
    nofo: UploadFile = File(...),
    rubric: Optional[UploadFile] = File(None),
    worksheet: Optional[UploadFile] = File(None),
    approved_criteria: str = Form(...),          # JSON string
    agency: str = Form("HRSA"),
    user_id: str = Form(...),
    review_id: str = Form(...),
):
    """
    Accept uploaded files, store them to Supabase Storage, create DB records,
    and enqueue per-application scoring as background tasks.
    Returns immediately with job IDs.
    """
    sb = get_supabase()

    # -- Parse approved_criteria --
    try:
        criteria: list[dict[str, Any]] = json.loads(approved_criteria)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=422, detail=f"approved_criteria is not valid JSON: {exc}")

    prefix = f"{user_id}/{review_id}"

    # -- Upload NOFO --
    nofo_bytes = await nofo.read()
    nofo_filename = nofo.filename or "nofo.pdf"
    nofo_storage_path = f"{prefix}/nofo/{nofo_filename}"
    _upload_bytes(sb, BUCKET_NOFO, nofo_storage_path, nofo_bytes, "application/pdf")

    # -- Upload worksheet template (optional) --
    worksheet_storage_path: Optional[str] = None
    if worksheet and worksheet.filename:
        ws_bytes = await worksheet.read()
        ws_filename = worksheet.filename
        worksheet_storage_path = f"{prefix}/worksheet/{ws_filename}"
        _upload_bytes(sb, BUCKET_WORKSHEETS, worksheet_storage_path, ws_bytes,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    # -- Process each application (may include ZIPs) --
    job_records: list[dict[str, Any]] = []

    for upload in applications:
        app_bytes = await upload.read()
        original_name = upload.filename or "application.pdf"
        suffix = Path(original_name).suffix.lower()

        if suffix == ".zip":
            # Extract ZIP in a temp dir, then handle each PDF
            with tempfile.TemporaryDirectory() as tmpdir:
                zip_path = Path(tmpdir) / original_name
                zip_path.write_bytes(app_bytes)
                extract_dir = Path(tmpdir) / "extracted"
                try:
                    _, safe_extract_application_zip, *_ = _lazy_scoring()
                    pdf_paths = safe_extract_application_zip(zip_path, extract_dir)
                except ValueError as exc:
                    raise HTTPException(status_code=422, detail=f"ZIP error in {original_name}: {exc}")

                for pdf_path in pdf_paths:
                    job_rec = _create_application_job(
                        sb, pdf_path.read_bytes(), pdf_path.name,
                        prefix, criteria, agency, review_id, user_id,
                        nofo_storage_path, worksheet_storage_path,
                    )
                    job_records.append(job_rec)
                    background_tasks.add_task(
                        _process_job,
                        job_rec["job_id"],
                        job_rec["application_storage_path"],
                        criteria, agency, review_id,
                        nofo_storage_path, worksheet_storage_path,
                    )
        else:
            # Single PDF
            job_rec = _create_application_job(
                sb, app_bytes, original_name,
                prefix, criteria, agency, review_id, user_id,
                nofo_storage_path, worksheet_storage_path,
            )
            job_records.append(job_rec)
            background_tasks.add_task(
                _process_job,
                job_rec["job_id"],
                job_rec["application_storage_path"],
                criteria, agency, review_id,
                nofo_storage_path, worksheet_storage_path,
            )

    return JSONResponse(content={
        "review_id": review_id,
        "jobs": job_records,
        "status": "queued",
        "message": f"{len(job_records)} application(s) queued for processing.",
    })


def _create_application_job(
    sb: Client,
    pdf_bytes: bytes,
    filename: str,
    prefix: str,
    criteria: list[dict[str, Any]],
    agency: str,
    review_id: str,
    user_id: str,
    nofo_storage_path: str,
    worksheet_storage_path: Optional[str],
) -> dict[str, Any]:
    """Upload application PDF, insert application + processing_job rows, return job metadata."""
    application_id = str(uuid.uuid4())
    job_id = str(uuid.uuid4())
    app_storage_path = f"{prefix}/applications/{application_id}_{filename}"

    _upload_bytes(sb, BUCKET_APPS, app_storage_path, pdf_bytes, "application/pdf")

    # Insert application record
    _insert(sb, "applications", {
        "id": application_id,
        "review_id": review_id,
        "user_id": user_id,
        "filename": filename,
        "storage_path": app_storage_path,
        "status": "queued",
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    # Insert processing_job record
    _insert(sb, "processing_jobs", {
        "id": job_id,
        "review_id": review_id,
        "application_id": application_id,
        "user_id": user_id,
        "status": "queued",
        "agency": agency,
        "criteria": json.dumps(criteria),
        "nofo_storage_path": nofo_storage_path,
        "worksheet_storage_path": worksheet_storage_path,
        "created_at": datetime.now(timezone.utc).isoformat(),
    })

    return {
        "job_id": job_id,
        "application_id": application_id,
        "filename": filename,
        "application_storage_path": app_storage_path,
        "status": "queued",
    }


# ---------------------------------------------------------------------------
# Pre-scoring document audit
# ---------------------------------------------------------------------------

class DocumentMismatchError(Exception):
    """Raised when NOFO, application, and worksheet funding opportunity numbers don't match."""
    pass


def _extract_funding_opportunity_number(pdf_bytes: bytes, label: str) -> str | None:
    """Extract Funding Opportunity Number from the first few pages of a PDF."""
    import fitz
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # Check first 5 pages for the FON
        text = ""
        for i in range(min(5, len(doc))):
            text += doc[i].get_text() + "\n"
        doc.close()

        # Common patterns for Funding Opportunity Number
        patterns = [
            r'(?:Funding\s+Opportunity\s+Number|NOFO\s+Number|FON|Opportunity\s+Number)\s*[:\-]?\s*([\w\-]+\-\d+\-\d+)',
            r'(HRSA-\d{2}-\d{3})',
            r'(SM-\d{2}-\d{3})',
            r'(TI-\d{2}-\d{3})',
            r'(SP-\d{2}-\d{3})',
            r'(CDC-\d{4}-\d+)',
            r'(NIH-\w+-\d{2}-\d{3})',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    except Exception as exc:
        logger.warning("Could not extract FON from %s: %s", label, exc)
        return None


def _extract_fon_from_worksheet(ws_bytes: bytes) -> str | None:
    """Extract Funding Opportunity Number from a DOCX worksheet."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(ws_bytes))
        text = "\n".join(p.text for p in doc.paragraphs[:30])
        # Also check tables (HRSA worksheets often use tables)
        for table in doc.tables[:5]:
            for row in table.rows:
                for cell in row.cells:
                    text += "\n" + cell.text

        patterns = [
            r'(?:Funding\s+Opportunity\s+Number|NOFO\s+Number|FON|Grant\s+Number)\s*[:\-]?\s*([\w\-]+\-\d+\-\d+)',
            r'(HRSA-\d{2}-\d{3})',
            r'(SM-\d{2}-\d{3})',
            r'(TI-\d{2}-\d{3})',
            r'(SP-\d{2}-\d{3})',
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return None
    except Exception as exc:
        logger.warning("Could not extract FON from worksheet: %s", exc)
        return None


def _prescore_document_audit(
    sb: Any,
    job_id: str,
    application_id: str,
    review_id: str,
    app_bytes: bytes,
    nofo_bytes: bytes | None,
    worksheet_storage_path: str | None,
) -> None:
    """Verify NOFO, application, and worksheet all reference the same Funding Opportunity Number.

    Raises DocumentMismatchError if they clearly don't match.
    Logs warnings for soft mismatches (e.g., one document missing the FON).
    """
    fon_nofo = _extract_funding_opportunity_number(nofo_bytes, "NOFO") if nofo_bytes else None
    fon_app = _extract_funding_opportunity_number(app_bytes, "Application")

    fon_worksheet = None
    if worksheet_storage_path:
        try:
            ws_bytes = _download_bytes(sb, BUCKET_WORKSHEETS, worksheet_storage_path)
            fon_worksheet = _extract_fon_from_worksheet(ws_bytes)
        except Exception as exc:
            logger.warning("Could not download worksheet for audit: %s", exc)

    logger.info("Document audit — NOFO FON: %s, Application FON: %s, Worksheet FON: %s",
                fon_nofo, fon_app, fon_worksheet)

    # Collect all non-None FONs
    fons = {}
    if fon_nofo:
        fons["NOFO"] = fon_nofo.upper()
    if fon_app:
        fons["Application"] = fon_app.upper()
    if fon_worksheet:
        fons["Worksheet"] = fon_worksheet.upper()

    if len(fons) < 2:
        # Can't compare if we only have one or zero FONs
        if not fons:
            logger.warning("Could not extract Funding Opportunity Number from any document — skipping audit")
        else:
            logger.info("Only extracted FON from %s (%s) — cannot cross-verify", list(fons.keys())[0], list(fons.values())[0])
        return

    # Check if all extracted FONs match
    unique_fons = set(fons.values())
    if len(unique_fons) == 1:
        logger.info("Document audit PASSED — all documents reference %s", list(unique_fons)[0])
        return

    # Mismatch detected — build error message
    details = ", ".join(f"{doc}: {fon}" for doc, fon in fons.items())
    raise DocumentMismatchError(
        f"Funding Opportunity Number mismatch across documents. {details}. "
        "Please ensure the NOFO, application, and reviewer worksheet all reference the same grant opportunity."
    )


# ---------------------------------------------------------------------------
# Background: _process_job
# ---------------------------------------------------------------------------

def _process_job(
    job_id: str,
    application_storage_path: str,
    criteria: list[dict[str, Any]],
    agency: str,
    review_id: str,
    nofo_storage_path: str,
    worksheet_storage_path: Optional[str],
) -> None:
    """
    Background task: download application from Storage, score with Claude,
    write results to DB, optionally populate worksheet, mark job complete.
    """
    sb = get_supabase()
    logger.info("Processing job %s", job_id)

    _update(sb, "processing_jobs", {"id": job_id}, {
        "status": "processing",
        "started_at": datetime.now(timezone.utc).isoformat(),
    })

    try:
        # -- Resolve application_id and storage path from DB --
        job_rows = _select(sb, "processing_jobs", {"id": job_id})
        if not job_rows:
            raise RuntimeError(f"Job {job_id} not found in DB")
        job_row = job_rows[0]
        application_id: str = job_row["application_id"]

        # Get actual storage path from applications table (application_storage_path arg may be UUID from retry)
        app_rows = _select(sb, "applications", {"id": application_id})
        actual_app_path = app_rows[0]["storage_path"] if app_rows else application_storage_path
        if not actual_app_path or len(actual_app_path) < 10:
            raise RuntimeError(f"No valid storage path for application {application_id}")

        # Also resolve nofo/worksheet paths from the application record if not passed
        if not nofo_storage_path and app_rows:
            nofo_storage_path = app_rows[0].get("nofo_storage_path", "")
        if not worksheet_storage_path and app_rows:
            worksheet_storage_path = app_rows[0].get("worksheet_storage_path", "")

        # -- Download application PDF --
        app_bytes = _download_bytes(sb, BUCKET_APPS, actual_app_path)

        # -- Download NOFO for guidance text --
        guidance_text = ""
        nofo_bytes = None
        try:
            nofo_bytes = _download_bytes(sb, BUCKET_NOFO, nofo_storage_path)
            try:
                from scoring.document_processor import DocumentProcessor
                proc = DocumentProcessor()
            except ImportError:
                proc = None
            if proc is None:
                pass  # skip guidance extraction
            else:
                proc = DocumentProcessor()
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as ntmp:
                ntmp.write(nofo_bytes)
                ntmp_path = Path(ntmp.name)
            try:
                doc_result = proc.process_document(str(ntmp_path))
                guidance_text = doc_result.get("text_content", "")[:30000]
            finally:
                ntmp_path.unlink(missing_ok=True)
        except Exception as exc:
            logger.warning("Could not extract NOFO guidance text: %s", exc)

        # -- Pre-scoring audit: verify NOFO / Application / Worksheet match --
        _update(sb, "processing_jobs", {"id": job_id}, {
            "error_message": "Verifying NOFO / Application / Worksheet match...",
        })
        try:
            _prescore_document_audit(sb, job_id, application_id, review_id,
                                     app_bytes, nofo_bytes,
                                     worksheet_storage_path)
        except DocumentMismatchError as dme:
            logger.error("Document mismatch for job %s: %s", job_id, dme)
            _update(sb, "processing_jobs", {"id": job_id}, {
                "status": "failed",
                "error_message": f"Document mismatch: {dme}",
                "completed_at": datetime.now(timezone.utc).isoformat(),
            })
            _update(sb, "applications", {"id": application_id}, {
                "status": "failed",
                "updated_at": datetime.now(timezone.utc).isoformat(),
            })
            return
        except Exception as audit_exc:
            logger.warning("Pre-scoring audit warning (non-blocking): %s", audit_exc)

        _update(sb, "processing_jobs", {"id": job_id}, {
            "error_message": "Document audit passed. Scoring with Claude...",
        })

        # -- Score with Claude --
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as atmp:
            atmp.write(app_bytes)
            app_tmp_path = Path(atmp.name)

        try:
            _, _, score_application_with_claude, _ = _lazy_scoring()
            review_result = score_application_with_claude(
                application=app_tmp_path,
                criteria=criteria,
                agency=agency,
                guidance=guidance_text,
            )
        finally:
            app_tmp_path.unlink(missing_ok=True)

        # -- Write criterion_scores rows --
        for crit in review_result.get("criteria", []):
            _insert(sb, "criterion_scores", {
                "id": str(uuid.uuid4()),
                "review_id": review_id,
                "application_id": application_id,
                "job_id": job_id,
                "criterion_name": crit["name"],
                "score": crit.get("score"),
                "maximum_points": crit.get("maximum_points"),
                "score_rationale": crit.get("score_rationale", ""),
                "strengths": json.dumps(crit.get("strengths", [])),
                "mets": json.dumps(crit.get("mets", [])),
                "weaknesses": json.dumps(crit.get("weaknesses", [])),
                "subcriteria": json.dumps(crit.get("subcriteria", [])),
                "created_at": datetime.now(timezone.utc).isoformat(),
            })

        # -- Write review_findings rows (flattened per finding) --
        for crit in review_result.get("criteria", []):
            for finding_type in ("strengths", "mets", "weaknesses"):
                for finding in crit.get(finding_type, []):
                    # Support both old "pages" and new "application_pages" field names
                    app_pages = finding.get("application_pages", finding.get("pages", []))
                    row: dict[str, Any] = {
                        "id": str(uuid.uuid4()),
                        "review_id": review_id,
                        "application_id": application_id,
                        "job_id": job_id,
                        "criterion_name": crit["name"],
                        "finding_type": finding_type,
                        "comment": finding.get("comment", ""),
                        "pages": json.dumps(app_pages),
                        "created_at": datetime.now(timezone.utc).isoformat(),
                    }
                    if finding_type == "weaknesses":
                        row["nofo_requirement"] = finding.get("nofo_requirement", "")
                        row["nofo_pages"] = json.dumps(finding.get("nofo_pages", []))
                        row["impact"] = finding.get("impact", "")
                    _insert(sb, "review_findings", row)

        # -- Populate reviewer worksheet (if template provided) --
        worksheet_doc_id: Optional[str] = None
        if worksheet_storage_path:
            _update(sb, "processing_jobs", {"id": job_id}, {
                "error_message": "Scoring complete. Generating reviewer worksheet...",
            })
            try:
                ws_bytes = _download_bytes(sb, BUCKET_WORKSHEETS, worksheet_storage_path)
                ws_suffix = Path(worksheet_storage_path).suffix or ".docx"

                with tempfile.NamedTemporaryFile(suffix=ws_suffix, delete=False) as wstmp:
                    wstmp.write(ws_bytes)
                    ws_tmp_path = Path(wstmp.name)

                # Use applicant name for the filename, fallback to application ID
                applicant = review_result.get("applicant_name", "").strip()
                safe_name = re.sub(r'[^\w\s\-]', '', applicant).strip().replace(' ', '_') if applicant else application_id
                out_filename = f"{safe_name}_Completed_Worksheet{ws_suffix}"
                with tempfile.NamedTemporaryFile(suffix=ws_suffix, delete=False) as outtmp:
                    out_tmp_path = Path(outtmp.name)

                try:
                    _, _, _, populate_reviewer_worksheet = _lazy_scoring()
                    populated_path = populate_reviewer_worksheet(
                        template=ws_tmp_path,
                        output=out_tmp_path,
                        review=review_result,
                    )
                    completed_bytes = populated_path.read_bytes()
                    completed_storage_path = f"{application_id}/{out_filename}"
                    _upload_bytes(
                        sb, BUCKET_COMPLETED, completed_storage_path, completed_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                    # Insert generated_documents record
                    worksheet_doc_id = str(uuid.uuid4())
                    _insert(sb, "generated_documents", {
                        "id": worksheet_doc_id,
                        "review_id": review_id,
                        "application_id": application_id,
                        "job_id": job_id,
                        "document_type": "completed_worksheet",
                        "storage_bucket": BUCKET_COMPLETED,
                        "storage_path": completed_storage_path,
                        "filename": out_filename,
                        "created_at": datetime.now(timezone.utc).isoformat(),
                    })
                finally:
                    ws_tmp_path.unlink(missing_ok=True)
                    out_tmp_path.unlink(missing_ok=True)

            except Exception as ws_exc:
                logger.warning("Worksheet population failed for job %s: %s", job_id, ws_exc)

        # -- Update application status --
        _update(sb, "applications", {"id": application_id}, {
            "status": "completed",
            "final_score": review_result.get("final_score"),
            "maximum_score": review_result.get("maximum_score"),
            "review_status": review_result.get("review_status"),
            "applicant_name": review_result.get("applicant_name"),
            "application_number": review_result.get("application_number"),
            "overall_summary": review_result.get("overall_summary"),
            "full_result": json.dumps(review_result),
            "updated_at": datetime.now(timezone.utc).isoformat(),
        })

        # -- Mark job complete --
        _update(sb, "processing_jobs", {"id": job_id}, {
            "status": "completed",
            "completed_at": datetime.now(timezone.utc).isoformat(),
            "worksheet_document_id": worksheet_doc_id,
            "error_message": None,
        })

        logger.info("Job %s completed — score %s/%s", job_id,
                    review_result.get("final_score"), review_result.get("maximum_score"))

    except Exception as exc:
        logger.exception("Job %s failed: %s", job_id, exc)
        _update(sb, "processing_jobs", {"id": job_id}, {
            "status": "failed",
            "error_message": str(exc)[:2000],
            "completed_at": datetime.now(timezone.utc).isoformat(),
        })
        # Also mark the application as failed
        try:
            job_rows = _select(sb, "processing_jobs", {"id": job_id})
            if job_rows:
                _update(sb, "applications", {"id": job_rows[0]["application_id"]}, {
                    "status": "failed",
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                })
        except Exception:
            pass


# ---------------------------------------------------------------------------
# POST /jobs/{job_id}/process  (internal trigger — idempotent re-run)
# ---------------------------------------------------------------------------

@app.post("/jobs/{job_id}/process")
async def trigger_process_job(job_id: str, background_tasks: BackgroundTasks):
    """
    Internal endpoint to (re-)trigger processing of a queued or failed job.
    Looks up job metadata from DB and enqueues the background task.
    """
    sb = get_supabase()
    job_rows = _select(sb, "processing_jobs", {"id": job_id})
    if not job_rows:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")
    job_row = job_rows[0]

    if job_row["status"] == "processing":
        return {"message": "Job is already processing", "job_id": job_id}

    criteria = json.loads(job_row.get("criteria") or "[]")

    # Look up actual storage path from applications table
    app_rows = _select(sb, "applications", {"id": job_row["application_id"]})
    app_storage_path = app_rows[0]["storage_path"] if app_rows else ""

    import threading
    def _run_in_thread():
        try:
            _process_job(
                job_id, app_storage_path, criteria,
                job_row.get("agency", "HRSA"),
                job_row["review_id"],
                job_row.get("nofo_storage_path", ""),
                job_row.get("worksheet_storage_path", ""),
            )
        except Exception as exc:
            logger.error("Thread job %s failed: %s", job_id, exc)
            try:
                sb2 = get_supabase()
                _update(sb2, "processing_jobs", {"id": job_id}, {"status": "failed", "error_message": str(exc)[:2000]})
            except Exception:
                pass
    threading.Thread(target=_run_in_thread, daemon=True).start()
    return {"message": "Job re-queued", "job_id": job_id}


# ---------------------------------------------------------------------------
# GET /jobs/{job_id}/status
# ---------------------------------------------------------------------------

@app.get("/jobs/{job_id}/status")
def get_job_status(job_id: str):
    sb = get_supabase()
    rows = _select(sb, "processing_jobs", {"id": job_id})
    if not rows:
        raise HTTPException(status_code=404, detail="Job not found")
    row = rows[0]
    return {
        "job_id": job_id,
        "application_id": row.get("application_id"),
        "review_id": row.get("review_id"),
        "status": row.get("status"),
        "error_message": row.get("error_message"),
        "created_at": row.get("created_at"),
        "started_at": row.get("started_at"),
        "completed_at": row.get("completed_at"),
    }


# ---------------------------------------------------------------------------
# GET /reviews/{review_id}/results
# ---------------------------------------------------------------------------

@app.get("/reviews/{review_id}/results")
def get_review_results(review_id: str):
    """Return full scored results for all applications in a review."""
    sb = get_supabase()

    applications = _select(sb, "applications", {"review_id": review_id})
    if not applications:
        raise HTTPException(status_code=404, detail="Review not found or has no applications")

    results = []
    for app_row in applications:
        app_id = app_row["id"]
        scores = _select(sb, "criterion_scores", {"application_id": app_id})
        findings = _select(sb, "review_findings", {"application_id": app_id})
        docs = _select(sb, "generated_documents", {"application_id": app_id})

        criteria_out = []
        for s in scores:
            criteria_out.append({
                "criterion_name": s["criterion_name"],
                "score": s["score"],
                "maximum_points": s["maximum_points"],
                "score_rationale": s.get("score_rationale", ""),
                "strengths": json.loads(s.get("strengths") or "[]"),
                "mets": json.loads(s.get("mets") or "[]"),
                "weaknesses": json.loads(s.get("weaknesses") or "[]"),
                "subcriteria": json.loads(s.get("subcriteria") or "[]"),
            })

        results.append({
            "application_id": app_id,
            "filename": app_row.get("filename"),
            "status": app_row.get("status"),
            "final_score": app_row.get("final_score"),
            "maximum_score": app_row.get("maximum_score"),
            "review_status": app_row.get("review_status"),
            "applicant_name": app_row.get("applicant_name"),
            "application_number": app_row.get("application_number"),
            "overall_summary": app_row.get("overall_summary"),
            "criteria": criteria_out,
            "findings_count": len(findings),
            "documents": [
                {
                    "document_id": d["id"],
                    "document_type": d["document_type"],
                    "filename": d["filename"],
                }
                for d in docs
            ],
        })

    return {"review_id": review_id, "applications": results}


# ---------------------------------------------------------------------------
# DELETE /reviews/{review_id}
# ---------------------------------------------------------------------------

@app.delete("/reviews/{review_id}")
def delete_review(review_id: str):
    """Delete review, all associated DB records, and all storage files."""
    sb = get_supabase()

    # Gather application IDs first for storage cleanup
    applications = _select(sb, "applications", {"review_id": review_id})
    app_ids = [a["id"] for a in applications]

    # Delete generated documents + storage
    for app_id in app_ids:
        docs = _select(sb, "generated_documents", {"application_id": app_id})
        for doc in docs:
            try:
                sb.storage.from_(doc["storage_bucket"]).remove([doc["storage_path"]])
            except Exception as exc:
                logger.warning("Could not delete storage object %s: %s", doc["storage_path"], exc)
        sb.table("generated_documents").delete().eq("application_id", app_id).execute()

    # Delete review_findings, criterion_scores, processing_jobs, applications
    for app_id in app_ids:
        sb.table("review_findings").delete().eq("application_id", app_id).execute()
        sb.table("criterion_scores").delete().eq("application_id", app_id).execute()

    sb.table("processing_jobs").delete().eq("review_id", review_id).execute()
    sb.table("applications").delete().eq("review_id", review_id).execute()

    # Delete storage prefixes (best-effort, iterate by user_id prefix)
    # Attempt deletion by review_id suffix across buckets
    for bucket in (BUCKET_APPS, BUCKET_NOFO, BUCKET_WORKSHEETS):
        for app_row in applications:
            user_id = app_row.get("user_id", "")
            _delete_storage_prefix(sb, bucket, f"{user_id}/{review_id}")

    return {"deleted": True, "review_id": review_id, "applications_deleted": len(app_ids)}


# ---------------------------------------------------------------------------
# GET /reviews/{review_id}/worksheet/{application_id}
# ---------------------------------------------------------------------------

class DeleteApplicantDataRequest(BaseModel):
    review_id: str = ""
    confirmation: str = ""


@app.post("/reviews/{review_id}/delete-applicant-data")
def delete_applicant_data(review_id: str, body: DeleteApplicantDataRequest):
    """
    Permanently delete all applicant data for a review:
    - Storage: grant-applications bucket, completed-worksheets bucket
    - DB: generated_documents, review_findings, criterion_scores, processing_jobs, applications
    Preserves: grant_reviews row, nofo-files bucket, nofo_briefs, worksheet-templates.
    Requires confirmation == 'DELETE APPLICANT DATA'.
    """
    if body.confirmation != "DELETE APPLICANT DATA":
        raise HTTPException(
            status_code=400,
            detail="Confirmation phrase must be exactly 'DELETE APPLICANT DATA'",
        )

    sb = get_supabase()

    # Gather all applications for this review
    applications = _select(sb, "applications", {"review_id": review_id})
    app_ids = [a["id"] for a in applications]
    storage_objects_deleted = 0

    for app in applications:
        app_id = app["id"]

        # Delete from grant-applications bucket by storage_path
        if app.get("storage_path"):
            try:
                sb.storage.from_(BUCKET_APPS).remove([app["storage_path"]])
                storage_objects_deleted += 1
            except Exception as exc:
                logger.warning("Could not delete app storage %s: %s", app["storage_path"], exc)

        # Delete completed-worksheets bucket objects via generated_documents
        docs = _select(sb, "generated_documents", {"application_id": app_id})
        for doc in docs:
            bucket = doc.get("storage_bucket") or BUCKET_COMPLETED
            if doc.get("storage_path"):
                try:
                    sb.storage.from_(bucket).remove([doc["storage_path"]])
                    storage_objects_deleted += 1
                except Exception as exc:
                    logger.warning("Could not delete doc storage %s: %s", doc["storage_path"], exc)

        # Delete DB child records
        sb.table("generated_documents").delete().eq("application_id", app_id).execute()
        sb.table("review_findings").delete().eq("application_id", app_id).execute()
        sb.table("criterion_scores").delete().eq("application_id", app_id).execute()

    # Delete processing_jobs at review level
    sb.table("processing_jobs").delete().eq("review_id", review_id).execute()

    # Delete applications rows
    sb.table("applications").delete().eq("review_id", review_id).execute()

    # Update grant_reviews status
    _update(sb, "grant_reviews", {"id": review_id}, {
        "status": "applicant_data_deleted",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    })

    logger.info(
        "Applicant data deleted for review %s: %d applications, %d storage objects",
        review_id, len(app_ids), storage_objects_deleted,
    )

    return {
        "review_id": review_id,
        "applications_deleted": len(app_ids),
        "storage_objects_deleted": storage_objects_deleted,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


@app.get("/reviews/{review_id}/application/{application_id}/view-url")
def get_application_view_url(review_id: str, application_id: str):
    """Generate a 60-minute signed URL for viewing an application PDF."""
    sb = get_supabase()
    app_rows = (
        sb.table("applications")
        .select("storage_path, filename")
        .eq("id", application_id)
        .eq("review_id", review_id)
        .execute()
        .data
    )
    if not app_rows:
        raise HTTPException(status_code=404, detail="Application not found")
    app_row = app_rows[0]
    storage_path = app_row.get("storage_path")
    if not storage_path:
        raise HTTPException(status_code=404, detail="Application has no storage path")
    try:
        signed = _signed_url(sb, BUCKET_APPS, storage_path, expires_in=3600)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not generate view URL: {exc}")
    return {
        "url": signed,
        "expires_in": 3600,
        "filename": app_row.get("filename", "application.pdf"),
    }


@app.get("/reviews/{review_id}/worksheet/{application_id}")
def get_worksheet_download_url(review_id: str, application_id: str):
    """Generate a signed 60-minute download URL for the completed worksheet."""
    sb = get_supabase()

    docs = (
        sb.table("generated_documents")
        .select("*")
        .eq("application_id", application_id)
        .eq("review_id", review_id)
        .eq("document_type", "completed_worksheet")
        .execute()
        .data
    )
    if not docs:
        raise HTTPException(
            status_code=404,
            detail="No completed worksheet found for this application",
        )

    doc = docs[0]
    try:
        signed = _signed_url(sb, doc["storage_bucket"], doc["storage_path"], expires_in=3600)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Could not generate download URL: {exc}")

    return {
        "document_id": doc["id"],
        "filename": doc["filename"],
        "download_url": signed,
        "expires_in_seconds": 3600,
    }
