"""
Generate WEST VIRGINIA UNIVERSITY RESEARCH CORPORATION - Review Results.docx
from the full_result JSON stored in Supabase.
"""

import json
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Raw data (parsed from Supabase full_result) ─────────────────────────────

DATA = {
    "applicant_name": "West Virginia University Research Corporation",
    "application_number": "248134",
    "agency": "HRSA",
    "overview": {
        "applicant_information": "West Virginia University Research Corporation, Morgantown, WV, is applying for continuation of the Leadership Education in Neurodevelopmental and Other Related Disabilities (WV LEND) program at Tier 3 funding level. The program is housed within the WVU Center for Excellence in Disabilities (WVUCED). Dr. Cassaundra Miller serves as Project Director/Principal Investigator (30% effort, EdD in Educational Leadership) and Dr. Margaret Jaynes serves as Co-Project Director (15% effort, MD, Professor Emeritus of Pediatrics & Neurology). The application requests $625,160 for Year 1 and $3,132,840 total over five years (pages 4-5, 16-32).",
        "target_population": "The target population is children and youth with autism and other developmental disabilities (autism/DD) across their lifespan and their families in West Virginia's predominantly rural communities. WV has the second highest percentage of children with special health care needs in the nation at 33.8% compared to 26.2% nationally. The state has approximately 0.3 developmental pediatricians per 100,000 children compared to a national average of 1:100,000. Almost 50% of West Virginians live in rural areas, more than double the national average of 20%. One in three West Virginians earn below 150% of the federal poverty level, and 25% of WV families live in food deserts. The state has a disability prevalence rate of 19.5% compared to 13.4% nationally (pages 2-6).",
        "project_description": "WV LEND will train 18 long-term trainees from 10 disciplines and 15 medium-term trainees annually through an interdisciplinary curriculum combining academic, clinical, leadership, and community-based experiences. The program utilizes ten curriculum forums: Introductory Disability Institute, Academic Content, Family Engagement, Leadership Seminar, Clinical, Community, Disability Identity Badge and Project, Research, Journal Club, and Self Reflection. Clinical experiences include the WVUCED Feeding and Swallowing Clinic, LEND Support Program, WVU Medicine Golisano Children's Neurodevelopmental Center, and community-based placements. The program will provide at least three continuing education activities annually and technical assistance to Title V programs, community organizations, and practicing professionals. WV LEND participates in the LEND Rural Collaborative with three other rural LEND programs (pages 1, 6-23).",
        "goals_objectives": "Goal 1: Provide evidence-based interdisciplinary education and training to advance understanding of children/youth with autism/DD by increasing academic/clinical knowledge and building leadership skills (10 objectives). Goal 2: Increase identification and use of evidence-based screening, diagnostics, and treatment for children/youth with autism/DD through exemplary, family-centered, interdisciplinary training (7 objectives). Goal 3: Develop and enhance innovative practice models that strengthen systems of care for individuals with autism/DD in rural communities through partnerships (6 objectives). Goal 4: Increase expertise of trainees and practicing professionals to recognize complex lifespan needs of children/youth with autism/DD with focus on lived experience perspective (6 objectives). The Work Plan in Attachment 1 details activities, responsibilities, evaluation methods, and timelines for each objective (pages 7-9, 26).",
        "significant_findings": "The applicant conducted a needs assessment with 15 stakeholders representing professionals, mentors, community members, and individuals with lived disability experience. All proposed MCH leadership competencies were rated as very important or extremely important. Highest rated competencies included working through problems, using a team approach, and considering individuals, environments, and resources when building services. Open-ended responses emphasized need for person-centered and trauma-informed care, evidence-based practices, community involvement, resource navigation, and accessible language. WV faces critical workforce shortages with only 2 developmental pediatricians statewide, resulting in wait times exceeding 6 months to 2+ years for evaluation. Medical mistrust is prevalent with over 25% of Appalachian respondents reporting high levels of healthcare mistrust. Three in five WV adults exhibit reading proficiency below 8th grade level (pages 2-6).",
        "other_information": "WV LEND has 25 years of MCH leadership and training experience. The program received over 100 applications for the 2026-2027 cohort. In 2025-26, twelve of nineteen trainees identified as having a disability and/or being a family member of an individual with a disability. The program is part of the national AUCD network of 68 UCEDDs. WV LEND maintains partnerships with WV Title V Programs (CSHCN, Birth to Three), WVU Medicine, WVU academic programs, Disability Action Center, and the LEND Rural Collaborative. The WVUCED Feeding and Swallowing Clinic is one of only two multidisciplinary feeding clinics in WV. The program uses PDSA Model for continuous quality improvement and collects data through WVUCED CODA database and AUCD NIRS system. Three former LEND trainees now serve as program mentors (pages 1, 10-12, 24-25, 31-35)."
    },
    "criteria": [
        {
            "name": "Need",
            "score": 9,
            "maximum_points": 10,
            "score_rationale": "The application presents compelling evidence of need through comprehensive data on WV's rural population, workforce shortages, and barriers to care. The needs assessment with 15 stakeholders demonstrates strong consensus on training priorities. Extensive state-specific data on poverty, food insecurity, disability prevalence, and healthcare access clearly establishes the context. The application would be strengthened by more specific data on current diagnostic evaluation wait times and quantified gaps in specific provider disciplines beyond developmental pediatricians.",
            "strengths": [
                {"comment": "The application presents robust state-specific data demonstrating WV has the second highest percentage of children with special health care needs nationally at 33.8% compared to 26.2%, and disability prevalence of 19.5% versus 13.4% nationally, clearly establishing the magnitude of need for trained providers serving children/youth with autism/DD.", "pages": [3, 4]},
                {"comment": "The applicant documents critical workforce shortages with only 2 developmental pediatricians statewide (0.3 per 100,000 children versus national average of 1:100,000), resulting in evaluation wait times exceeding 6 months and sometimes extending past 2 years, demonstrating urgent need for workforce development.", "pages": [4]},
                {"comment": "The needs assessment with 15 diverse stakeholders including individuals with lived experience, mentors, and community members shows all proposed MCH leadership competencies rated as very important or extremely important, with highest ratings for working through problems, team approach, and considering individuals/environments/resources when building services, providing strong stakeholder validation of training priorities.", "pages": [2, 3]},
                {"comment": "The application effectively describes multiple intersecting barriers facing WV families including geographic isolation (50% rural versus 20% national average), poverty (1 in 3 earning below 150% federal poverty level), food insecurity (25% in food deserts versus 5.6% nationally), limited transportation infrastructure (60% of roads rated fair to poor), and low health literacy (3 in 5 adults reading below 8th grade level), demonstrating comprehensive understanding of contextual factors affecting the target population.", "pages": [3, 4, 5]},
                {"comment": "The application cites recent peer-reviewed literature (Lindsay et al. 2026, Bican et al. 2025, Thomas et al. 2023, Efird et al. 2021) documenting rural-specific challenges including lack of local specialists, long waitlists, large travel distances (average 1.5 hours), diagnostic delays (average 2-year gap between parent concern and diagnosis), and medical mistrust (over 25% reporting high levels), strengthening the evidence base for need.", "pages": [4, 5]}
            ],
            "mets": [],
            "weaknesses": [
                {"comment": "While the application states evaluation wait times exceed 6 months and sometimes extend past 2 years, it does not provide specific quantified data on average wait times, number of children currently waiting for evaluation, or geographic distribution of wait times across the state, limiting ability to fully assess the scope of access barriers.", "pages": [4]}
            ]
        },
        {
            "name": "Response to Statement of Need",
            "score": 32,
            "maximum_points": 35,
            "score_rationale": "The application presents a comprehensive, well-designed training program with strong interdisciplinary curriculum, diverse clinical experiences, and clear alignment with program requirements. The ten curriculum forums, four-pillar framework (academic, clinical, community, leadership), and integration of MCH competencies demonstrate thoughtful program design. Recruitment strategies, mentor matching, and trainee support structures are well-articulated. The Work Plan provides detailed objectives with activities, responsibilities, and evaluation methods. Partnerships with WVUCED, WVU Medicine, Title V programs, and the LEND Rural Collaborative are substantive and well-documented. Minor weaknesses include limited detail on specific screening/diagnostic tools to be taught, incomplete description of medium-term trainee selection criteria, and lack of specific strategies for recruiting trainees from underrepresented populations beyond general statements about diversity.",
            "strengths": [
                {"comment": "The application proposes a comprehensive curriculum framework built on four foundational pillars (academic, clinical, community, leadership) delivered through ten curriculum forums including Introductory Disability Institute, Academic Content, Family Engagement, Leadership Seminar, Clinical, Community, Disability Identity Badge and Project, Research, Journal Club, and Self Reflection, demonstrating systematic integration of conceptual and performance training aligned with MCH principles.", "pages": [14, 15, 16]},
                {"comment": "The trainee recruitment strategy is multi-faceted including advertisements in student newspaper, announcements in Disability Studies courses, WVUCED/WVU/AUCD social media, emails from discipline mentors, university/community/statewide recruitment fairs, and referrals from current trainees, with collaborative relationships established with multiple WVU academic programs (Speech-Language Pathology, Physician Assistant Studies, School of Dentistry, Occupational Therapy) that actively promote LEND to students, demonstrating strong institutional support and diverse recruitment pathways.", "pages": [10, 11]},
                {"comment": "The application describes robust clinical training through multiple interdisciplinary clinics including WVUCED Feeding and Swallowing Clinic (one of only two multidisciplinary feeding clinics in WV), LEND Support Program, WVU Medicine Golisano Children's Neurodevelopmental Center, Cerebral Palsy and Spasticity Clinic, Complex Care Clinic, and CRADLE Lab, with trainees progressing from observation to direct participation and leadership roles under mentorship from recognized discipline leaders, providing extensive hands-on experience with the target population.", "pages": [20, 21, 22]},
                {"comment": "The program demonstrates strong commitment to including individuals with lived experience by requiring at least one trainee with a disability and one family member trainee in each cohort, employing dedicated family and disability mentors as core faculty, and reporting that twelve of nineteen trainees in 2025-26 identified as having a disability or being a family member of an individual with a disability, with three former LEND trainees now serving as lived experience mentors, creating a sustainable pipeline.", "pages": [11, 12, 16, 17]},
                {"comment": "The application presents substantive partnerships with clear roles and mutual benefits including WVUCED (housing organization providing infrastructure, programs, communications team), WVU Medicine/Golisano Children's Hospital (clinical training sites, mentor support), WV Title V Programs CSHCN and Birth to Three (technical assistance, trainee placement, service coordination), Disability Action Center (community-based experiential training), and LEND Rural Collaborative (shared journal club, continuing education, rural-focused curriculum development), with letters of support from WVU College of Law, Department of Pediatrics, and Physician Assistant Studies Program.", "pages": [24, 25, 26, 35]},
                {"comment": "The Work Plan in Attachment 1 provides detailed objectives organized by four goals with specific activities, person(s) responsible, evaluation methods, and timelines for each objective, demonstrating systematic planning and accountability structures to achieve program goals.", "pages": [7, 8, 9, 28]},
                {"comment": "The application describes a comprehensive approach to trainee support including pairing all trainees with primary mentors from same/similar disciplines, development of Individualized Training Plans with three goals (leadership, research/policy, personal development), dedicated trainee room with computers/internet/printing/lending library, access to all curriculum materials in multiple formats (synchronous in-person and virtual, asynchronous recorded, written and audio), and collaboration with WVU Office of Student Accommodations for specific needs.", "pages": [11, 12, 15, 16]},
                {"comment": "The continuing education plan includes at least three CE activities annually, established partnerships providing ongoing CE/TA to Monongalia County Schools (feeding needs in students with autism/DD), WV CSHCN (nutrition/feeding consultation), LinguaCare therapy company (monthly sessions plus office hours), WVUCED programs, and presentations at Mountain State Conference on Disabilities, with WVUCED serving as approved CE provider for social work, psychology, occupational therapy, speech therapy, physical therapy, and nursing.", "pages": [26, 27, 28]},
                {"comment": "The application demonstrates understanding of rural training needs by incorporating telehealth delivery for assessment and intervention, partnering with LEND Rural Collaborative for shared rural-focused journal club and CE development, including pre-staffing discussions of county demographics/school systems/community health factors, and providing training on plain language communication and accessible document creation to address lower literacy rates in Appalachian communities.", "pages": [5, 20, 25, 26]}
            ],
            "mets": [
                {"comment": "The application commits to recruiting 18 long-term trainees from 10 disciplines annually, ensuring no single discipline comprises more than 25% of the cohort, with at least one trainee with a disability and one family member trainee in each cohort, meeting NOFO requirements for Tier 3 trainee numbers and disciplines.", "pages": [7, 11]},
                {"comment": "The application commits to recruiting 15 medium-term trainees annually who will complete 40+ hours of defined curriculum including at least 19 clinical hours and 22 didactic hours, participating in core activities, interdisciplinary clinics, and LEND Journal Club alongside long-term trainees, and completing online LEND learning modules built as a Disability Micro-Credential Badge, meeting NOFO medium-term trainee requirements.", "pages": [7, 12, 13]},
                {"comment": "The application commits to identifying, recruiting and retaining mentors from 9 different disciplines per year with at least one lived experience and one family mentor, meeting NOFO requirements for Tier 3 faculty disciplines.", "pages": [7]},
                {"comment": "The curriculum covers all required topics including theoretical concepts (models of disability, critical thinking, research methods, interdisciplinary care, evidence-based practice, performance measurement, quality improvement), culturally and linguistically appropriate services, systems issues (systems of care for CYSHCN, person/family-centered care, medical home, care transitions, Title V MCH Services Block Grant), and clinical issues (autism screening/diagnosis/interventions, screening and diagnostic tools, social environment, transitions, services in adulthood, emerging needs, telehealth).", "pages": [14, 15, 16, 17, 18, 19, 20]},
                {"comment": "The application describes training delivery through balance of academic, clinical, and community opportunities including academic forums (Introductory Disability Institute, learning modules, trainee tracks, microlearning sessions, readings), clinical experiences (Feeding Clinic, LEND Support Program, NDC, multiple specialty clinics), and community placements (Disability Action Center, school-based activities, family engagement), meeting NOFO requirements for varied training settings.", "pages": [14, 15, 16, 17, 18, 19, 20, 21, 22, 23]},
                {"comment": "The application describes faculty qualifications including Project Director with EdD and over 15 years experience serving children with high intensity needs and as LEND Program Manager, Co-Project Director who is board-certified pediatric neurologist with 20 years as LEND director, and Clinical Coordinator who is MD with expertise in feeding disorders, along with core faculty from required disciplines who meet standards of education, experience, certification/licensure and demonstrate leadership, scholarship, and clinical experience with autism/DD population.", "pages": [9, 10, 11, 12, 13, 16, 17]}
            ],
            "weaknesses": [
                {"comment": "While the application states trainees will learn about screening tools and their functions and mentions STAT-MD training and ADOS-2 training, it does not provide sufficient detail on the specific screening and diagnostic tools that will be taught, the depth of training on each tool, or how competency in tool administration will be assessed, limiting ability to evaluate adequacy of diagnostic training.", "pages": [18, 19]},
                {"comment": "The application describes medium-term trainee recruitment from multiple sources (pediatric residents, child psychiatry fellows, practicum students, individuals seeking additional learning, non-selected long-term applicants) but does not clearly articulate selection criteria, application/interview process, or how the program will ensure medium-term trainees have appropriate background and commitment to serving children/youth with autism/DD.", "pages": [12, 13]},
                {"comment": "While the application makes general statements about recruiting trainees who represent various populations and describes WV's diverse challenges, it does not provide specific strategies for recruiting trainees from racial/ethnic minority backgrounds, rural communities, or other underrepresented groups beyond the strong focus on recruiting individuals with disabilities and family members.", "pages": [11, 13]}
            ]
        },
        {
            "name": "Performance Reporting and Evaluation",
            "score": 14,
            "maximum_points": 15,
            "score_rationale": "The application presents a comprehensive performance measurement and evaluation framework using multiple quantitative and qualitative methods aligned with program objectives. The PDSA Model for CQI is well-articulated with clear cycles for overall program, trainee development, and CE/TA activities. Data collection systems (CODA, NIRS, SOLE) are established and secure. Monthly data meetings with defined roles demonstrate commitment to ongoing monitoring and data-driven decision making. The evaluation plan addresses implementation fidelity, outcome measurement, and sustainability. Table 2 clearly links evaluation methods to program goals. Minor weaknesses include limited detail on specific benchmarks/targets for some objectives, incomplete description of how qualitative data will be systematically analyzed, and lack of specific timeline for when evaluation findings will be reviewed and acted upon beyond monthly meetings.",
            "strengths": [
                {"comment": "The application presents a comprehensive performance measurement framework using multiple methods including MCH trainee competency self-assessments, mentor evaluation rubrics, training participation records, clinic/seminar activity logs, ITPs, training evaluations, LEND follow-up surveys, exit interviews, and HRSA-required measures, with Table 2 clearly linking evaluation methods to specific program goals, demonstrating systematic approach to measuring all key program components.", "pages": [30, 31]},
                {"comment": "The application describes established data collection infrastructure using secure online systems including internal WVUCED tracking system (CODA) with two-step verification, AUCD web-based NIRS system for federal reporting compliance, and WV LEND Online Learning Platform (SOLE) for trainee data, with WVUCED Data Analyst responsible for establishing accounts, assisting with data entry, and generating customized reports, demonstrating robust data management capacity.", "pages": [31]},
                {"comment": "The CQI framework based on PDSA Model is well-articulated with clear cycles for overall program (Plan: develop goals/objectives based on needs; Do: implement competency-based curriculum; Study: collect process data through established systems; Act: measure impact through performance measures), trainee development (ITP development, monthly ISR review, mid-year evaluation with adjustments, end-of-year reflection), and CE/TA activities (training evaluations, attendance monitoring, qualitative feedback, licensure standards), demonstrating systematic approach to iterative improvement.", "pages": [33, 34]},
                {"comment": "The application commits to monthly data meetings with Program Director, Co-Program Director, Training Director, Program Manager, and WVUCED Data Analyst to review performance data, inform program improvements, and ensure implementation fidelity, with clear roles defined (PM/TD collect and organize trainee data, Program Administrator enters mentor data, Data Analyst prepares initial analysis and may create run charts for formative decisions), demonstrating regular monitoring and collaborative decision-making structures.", "pages": [30]},
                {"comment": "The evaluation plan addresses multiple purposes including monitoring implementation fidelity, measuring outcomes, providing timely data for CQI, collecting data for HRSA evaluation and Autism CARES Act reporting, providing direction for sustainability, developing shareable results, and informing federal reports to Congress, demonstrating comprehensive understanding of evaluation functions and stakeholder needs.", "pages": [32]},
                {"comment": "The application describes multimodal strategies for 5-year and 10-year post-training follow-up including web-based texting (identified through CQI as most effective method), email, phone calls, ongoing social media contact, media announcements via LEND/WVUCED webpages/social media, collection of permanent contact information during exit interviews, maintaining secure repository of trainee phone numbers, and placing former trainees on list-serve, with ongoing reminders throughout training about importance of maintaining contact.", "pages": [10, 31, 32]},
                {"comment": "The application demonstrates commitment to using evaluation findings for program improvement by describing how CQI findings are used to strengthen trainee experiences, enhance interdisciplinary education, identify gaps in trainee support, and ensure curriculum remains responsive to emerging needs, with CQI reporting occurring at least twice yearly to mentors, trainees, advisory committees, and leadership to promote accountability and transparency.", "pages": [34]}
            ],
            "mets": [
                {"comment": "The application commits to tracking and reporting required performance measures including number of long-, medium-, and short-term trainees by discipline, number of CE events and participants, number of TA activities and participants, number of diagnostic evaluations, details on technical assistance and subject matter expertise, number of individuals taught to diagnose/screen/provide interventions, details on impact on systems of care, and progress on supplements, meeting NOFO reporting requirements.", "pages": [18, 19, 30, 31]},
                {"comment": "The application describes participation in evaluation of Autism CARES Act awardees which may inform federal reports to Congress, meeting NOFO requirement for contributing to federal evaluation efforts.", "pages": [19, 32]}
            ],
            "weaknesses": [
                {"comment": "While the application describes evaluation methods and data collection systems, it does not provide specific quantitative benchmarks or targets for most objectives (e.g., target percentage of trainees demonstrating competency, target satisfaction ratings for CE activities, target number of diagnostic evaluations), limiting ability to assess whether proposed evaluation will adequately measure success.", "pages": [30, 31]},
                {"comment": "The application describes collection of qualitative data through exit interviews, follow-up surveys, and training evaluations but does not describe how qualitative data will be systematically analyzed (e.g., coding procedures, thematic analysis methods, inter-rater reliability), limiting confidence in rigor of qualitative evaluation component.", "pages": [30, 31]}
            ]
        },
        {
            "name": "Impact",
            "score": 9,
            "maximum_points": 10,
            "score_rationale": "The application demonstrates strong potential for impact through comprehensive training of 18 long-term and 15 medium-term trainees annually across 10 disciplines, extensive CE/TA activities reaching 200+ practicing professionals yearly, partnerships with Title V programs and community organizations, and focus on rural workforce development. The program's 25-year history, track record of training over 24,000 trainees in FY23 with 153,000+ diagnostic evaluations, and pipeline of former trainees becoming mentors demonstrate sustained impact. Integration within WVUCED infrastructure and partnerships with WVU Medicine, academic programs, and LEND Rural Collaborative position the program for broad reach. The application would be strengthened by more specific projections of number of diagnostic evaluations to be provided, quantified targets for workforce retention in WV, and clearer articulation of how impact will be measured at community and systems levels.",
            "strengths": [
                {"comment": "The application demonstrates potential for significant workforce impact by training 18 long-term trainees from 10 disciplines annually who will receive 300+ hours of interdisciplinary training in screening, evaluation, diagnosis, and evidence-based intervention for autism/DD, with 12 of 19 current trainees identifying as having disability/family experience, creating a diverse workforce equipped to serve rural populations and address critical shortage of providers (only 2 developmental pediatricians statewide).", "pages": [1, 4, 7, 11]},
                {"comment": "The program demonstrates track record of substantial impact with over 24,000 trainees trained and continuing education provided to 433,000 practicing professionals in FY23, and over 153,000 infants and children receiving diagnostic evaluations through LEND programs that year, indicating capacity for large-scale workforce development and service delivery.", "pages": [4]},
                {"comment": "The application describes broad dissemination reach through WVUCED network of approximately 5,000 individuals with disabilities, individuals with lived experience, policymakers, stakeholders, educators, current/past trainees, and community organizations, plus amplification through AUCD national platforms (conferences, webinars, publications, professional networks) and WVU statewide clinical/educational/community networks, ensuring program products and training reach local, state, and national audiences.", "pages": [6, 7]},
                {"comment": "The application demonstrates systems-level impact potential through partnerships with WV Title V programs (CSHCN, Birth to Three) providing TA/coaching, mentor service on state advisory boards (WV CSHCN Medical Advisory Board, WV Infant and Toddler Mental Health Association Board, Governor's Early Childhood Advisory Council), and development of innovative service delivery models (telehealth, interdisciplinary feeding clinic, LEND Support Program) that can be replicated to improve rural systems of care.", "pages": [25, 26, 27, 28]},
                {"comment": "The program demonstrates sustainable impact through pipeline development with three former LEND trainees now serving as program mentors, integration of LEND training into multiple WVU academic programs (with letters of support from College of Law, Department of Pediatrics, Physician Assistant Studies), expansion of WVU Nutrition graduate curriculum to include week-long LEND dietetic internship for all students, and 25 years of established MCH leadership training experience.", "pages": [11, 35]},
                {"comment": "The application describes impact on continuing education and technical assistance by committing to at least three CE activities annually, providing ongoing TA/coaching to Monongalia County Schools, WV CSHCN, LinguaCare therapy company, and WVUCED programs, presenting at Mountain State Conference on Disabilities, and reaching at least 200 practicing professionals yearly through consultation, TA, dissemination, and CE activities.", "pages": [7, 9, 26, 27, 28]}
            ],
            "mets": [
                {"comment": "The application commits to increasing by 10% yearly the number of long-term trainees reporting engagement in policy development, implementation and evaluation, demonstrating commitment to measurable impact on policy and systems change.", "pages": [8]},
                {"comment": "The application commits to increasing by 10% per year the percent of long-term trainees who at 5 and 10 years post-training work in an interdisciplinary manner to serve the MCH population and individuals with autism/DD and their families, demonstrating commitment to measuring long-term workforce impact.", "pages": [10]}
            ],
            "weaknesses": [
                {"comment": "While the application describes providing diagnostic evaluations through LEND faculty and trainees and cites 153,000+ evaluations nationally in FY23, it does not provide specific projections or targets for the number of diagnostic evaluations WV LEND will conduct annually during the project period, limiting ability to assess anticipated direct service impact.", "pages": [4, 8, 18]}
            ]
        },
        {
            "name": "Resources and Capabilities",
            "score": 19,
            "maximum_points": 20,
            "score_rationale": "The application demonstrates strong organizational capacity through integration within WVUCED (federally designated UCEDD with 25 years of LEND experience), access to WVU and WVU Medicine infrastructure, and established partnerships. The organizational structure with clearly defined roles (PD, Co-PD, Training Director, Program Manager, Data Analyst, Dissemination Officer) and staffing plan in Attachment 2 shows appropriate expertise and time commitments. Faculty qualifications are strong with PD having EdD and 15+ years disability experience, Co-PD being board-certified pediatric neurologist with 20 years LEND director experience, and core faculty representing required disciplines with appropriate credentials. The budget narrative demonstrates reasonable allocation of resources. Minor weaknesses include PD effort at 30% being at the minimum recommended level, limited detail on how the program will ensure adequate mentor time given competing clinical/academic demands, and incomplete description of specific facilities and equipment available for trainee use beyond the trainee room.",
            "strengths": [
                {"comment": "The application demonstrates strong organizational capacity by housing WV LEND within WVUCED, a federally designated University Center for Excellence in Developmental Disabilities (UCEDD) part of national network of 68 UCEDDs coordinated through AUCD, with established infrastructure including communications team (7% FTE Dissemination Officer funded by LEND), data management systems (CODA, NIRS), and multiple programs (Disability Studies, PBS, WVATS, Specialized Family Care, Feeding and Swallowing Clinic) that provide comprehensive services and training opportunities.", "pages": [24, 25]},
                {"comment": "The Project Director demonstrates appropriate qualifications with EdD in Educational Leadership (2021), over 15 years experience as pediatric speech-language pathologist serving children with high intensity needs including decade in early intervention, current role as Assistant Professor in Disability Studies and Clinic Director at WVUCED, experience as LEND Program Manager since 2021, and 30% effort dedicated to LEND (7.8 calendar months), meeting NOFO requirement for at least 30% PD effort.", "pages": [9, 10, 11]},
                {"comment": "The Co-Project Director demonstrates strong qualifications as board-certified pediatric neurologist, Professor Emeritus of Pediatrics & Neurology, 20 years experience as LEND Project Director, expertise in telehealth (served on WV Telehealth Alliance Board 2010-2017 as chairman 2014-2015, Advisory Board member to Mid Atlantic Regional Telehealth Center), and current Project Director for Paths for Parents program, with 15% effort (1.8 calendar months Year 1, 0.6 months Years 2-5) dedicated to LEND.", "pages": [13, 16]},
                {"comment": "The application describes qualified core faculty including Clinical Coordinator (MD with expertise in feeding disorders, 6.0 calendar months Year 1, 5.4 months Years 2-5), mentors from at least 9 disciplines meeting standards of education/experience/certification/licensure with demonstrated leadership and clinical experience serving children/youth with autism/DD, dedicated family faculty member and faculty member with disability, and supporting faculty to strengthen curriculum and provide mentorship, with organizational chart in Attachment 4 showing roles and relationships.", "pages": [16, 17]},
                {"comment": "The staffing plan demonstrates appropriate allocation of personnel including 12 other personnel positions covering communications/dissemination (2.4 calendar months), data (1.2 months), clinical staff (9.84 months Year 1, 8.64 months Years 2-5), mentors (5.28 months), finance (2.4 months), and secretarial/clerical support (24 months), with Attachment 2 providing job descriptions showing qualifications and responsibilities aligned with program needs.", "pages": [16, 17]},
                {"comment": "The application demonstrates access to extensive clinical training infrastructure through partnerships with WVU Medicine Golisano Children's Neurodevelopmental Center (developmental pediatrician mentor, diagnostic clinics, therapy services), WVUCED Feeding and Swallowing Clinic (one of only two multidisciplinary feeding clinics in WV), WVU Medicine specialty clinics (Cerebral Palsy and Spasticity, Complex Care), WVU Speech Language and Hearing Clinic, and CRADLE Lab, providing diverse high-quality clinical experiences.", "pages": [4, 20, 21, 22, 25]},
                {"comment": "The application demonstrates institutional commitment through cost-sharing with WVUCED for office space/rent/technology and faculty/staff salaries for accounting, communications, and data management, letters of support from WVU College of Law, Department of Pediatrics, and Physician Assistant Studies Program describing continued support and value of LEND training, and integration of LEND into academic curricula (WVU Nutrition requiring week-long LEND dietetic internship for all graduate students).", "pages": [35]}
            ],
            "mets": [
                {"comment": "The Project Director meets NOFO qualifications as health professional with doctorate (EdD) in approved discipline and at least 5 years experience in programs serving children with autism/DD (over 15 years experience), with 30% effort dedicated to LEND meeting minimum requirement.", "pages": [9, 10, 11, 16]},
                {"comment": "The application describes organization with mission and structure enabling successful implementation, including WVUCED mission to improve lives of West Virginians through evidence-based services, training, research, and information sharing, experience working with populations with limited access to services, ability to recruit trainees in proposed disciplines, and infrastructure ensuring accessibility of settings/materials/curricula for all participants.", "pages": [17, 24, 25]}
            ],
            "weaknesses": [
                {"comment": "While the Project Director's 30% effort (7.8 calendar months) meets the NOFO minimum requirement, this is at the lower end of recommended effort for managing a comprehensive training program with 18 long-term trainees, 15 medium-term trainees, multiple clinical sites, continuing education activities, and complex partnerships, raising questions about adequacy of PD time for day-to-day oversight and trainee mentorship.", "pages": [16]}
            ]
        },
        {
            "name": "Support Requested",
            "score": 10,
            "maximum_points": 10,
            "score_rationale": "The budget is reasonable, well-justified, and aligned with proposed activities. Year 1 request of $625,160 is within Tier 3 ceiling of $641,000. The budget allocates appropriate resources to personnel (salaries and fringe for PD, Co-PD, Clinical Coordinator, and 12 other personnel), participant/trainee support ($198,000 Year 1 for 18 trainees), travel, consultant services, and other direct costs. The budget narrative provides detailed justification for all line items. Indirect costs are correctly calculated at 8% of MTDC per training grant requirements. The five-year budget shows appropriate escalation and consistency across years. All costs appear necessary and allocable to program objectives.",
            "strengths": [
                {"comment": "The Year 1 budget request of $625,160 is reasonable and within the Tier 3 funding ceiling of $641,000, with total five-year request of $3,132,840 demonstrating appropriate resource planning for the period of performance.", "pages": [16, 32]},
                {"comment": "The budget allocates appropriate resources to trainee support with $198,000 in participant/trainee support costs for 18 trainees in Year 1 (average $11,000 per trainee), decreasing to $176,000 in Years 2-5, demonstrating reasonable stipend levels to support graduate/postgraduate trainees for 9-12 months in accordance with WVU policies.", "pages": [17, 20, 23, 26, 29, 12]},
                {"comment": "Personnel costs are reasonable with PD at 7.8 calendar months ($34,418 salary Year 1), Co-PD at 1.8 months ($25,650 Year 1 decreasing to 0.6 months/$11,400 Years 2-5), Clinical Coordinator at 6.0 months ($41,921 Year 1), and 12 other personnel totaling $62,478 Year 1, with appropriate fringe benefits calculated, demonstrating adequate staffing to support 18 long-term and 15 medium-term trainees plus CE/TA activities.", "pages": [16, 19, 22, 25, 28]},
                {"comment": "Other direct costs are well-justified including consultant services ($66,316 Year 1 for mentors from multiple disciplines), rent/utilities/telecommunications ($47,000 Year 1 for WVUCED space and technology), materials and supplies ($11,778 Year 1), and smaller amounts for printing/postage and hospitality, all aligned with proposed training activities and partnerships.", "pages": [18, 21, 24, 27, 30]},
                {"comment": "Indirect costs are correctly calculated at 8% of modified total direct costs ($28,308 on base of $353,852 Year 1) in accordance with 2 CFR 200.414 requirement that indirect costs for training awards cannot exceed 8% of MTDC, with cognizant federal agency identified as DHHS.", "pages": [18, 21, 24, 27, 30]},
                {"comment": "The budget demonstrates appropriate cost allocation across the five-year period with Year 1 total of $625,160 and Years 2-5 totals of $626,920 each, showing consistency in resource needs with minor adjustments for personnel changes (Co-PD effort reduction, salary increases) and activity modifications (reduced materials/supplies, travel adjustments), indicating thoughtful long-term planning.", "pages": [16, 19, 22, 25, 28, 32]}
            ],
            "mets": [
                {"comment": "The budget includes no equipment costs, appropriate for a training program, and no program income is anticipated, meeting NOFO requirements.", "pages": [17, 20, 23, 26, 29]},
                {"comment": "All personnel salaries respect the Executive Level II salary cap of $228,000 as of January 2026, with highest individual salary shown as $228,000 for Co-PD base salary, meeting federal salary limitation requirements.", "pages": [16, 19, 22, 25, 28]}
            ],
            "weaknesses": []
        }
    ],
    "budget": {
        "recommendation": "as_requested",
        "annual_recommended_funding": [625160, 626920, 626920, 626920, 626920],
        "reduction_rationale": "No budget reduction is recommended. The requested budget of $625,160 for Year 1 and $626,920 for Years 2-5 (total $3,132,840) is reasonable, well-justified, and aligned with proposed program activities. All costs are necessary and allocable to training 18 long-term and 15 medium-term trainees annually, providing continuing education to practicing professionals, conducting diagnostic evaluations, and supporting technical assistance activities. Personnel costs are appropriate for the proposed effort levels and number of staff needed to support comprehensive interdisciplinary training. Trainee support costs of $198,000 Year 1 and $176,000 Years 2-5 are reasonable for stipends and related support. Other direct costs including consultant services for mentors ($66,316 Year 1, $43,316 Years 2-5), rent/utilities ($47,000 Year 1, $46,000-$46,500 Years 2-5), and materials/supplies are justified by proposed activities. Indirect costs are correctly calculated at 8% of MTDC per training grant requirements. The budget is within the Tier 3 ceiling of $641,000 annually and demonstrates appropriate resource allocation to achieve program objectives."
    },
    "overall_summary": "This is a well-designed competing continuation application from West Virginia University Research Corporation for the WV LEND program at Tier 3 funding level. The application demonstrates compelling need based on WV's status as second-highest state for children with special health care needs (33.8% vs 26.2% nationally), critical shortage of developmental pediatricians (0.3 per 100,000 vs national average of 1:100,000), and multiple intersecting barriers facing rural Appalachian families including poverty, geographic isolation, limited transportation, and low health literacy. The needs assessment with 15 stakeholders validates training priorities with all MCH leadership competencies rated as very/extremely important.\n\nThe response is comprehensive and well-aligned with NOFO requirements. The program will train 18 long-term trainees from 10 disciplines and 15 medium-term trainees annually through a four-pillar curriculum framework (academic, clinical, community, leadership) delivered via ten forums. Clinical training through WVUCED Feeding and Swallowing Clinic, WVU Medicine Golisano Children's Neurodevelopmental Center, and multiple specialty clinics provides extensive hands-on interdisciplinary experience. Strong commitment to lived experience is demonstrated with at least one trainee with disability and one family member trainee required per cohort, dedicated lived experience mentors, and 12 of 19 current trainees identifying as having disability/family experience. Partnerships with WVUCED, WVU Medicine, Title V programs (CSHCN, Birth to Three), and LEND Rural Collaborative are substantive with clear mutual benefits. The Work Plan provides detailed objectives with activities, responsibilities, and evaluation methods.\n\nPerformance measurement and evaluation are strong with comprehensive framework using multiple methods (competency assessments, mentor evaluations, activity logs, ITPs, surveys, exit interviews), established data systems (CODA, NIRS, SOLE), monthly data meetings with defined roles, and well-articulated PDSA Model for CQI. The plan addresses implementation fidelity, outcome measurement, HRSA reporting, and sustainability.\n\nImpact potential is significant through training 33 trainees annually across 10 disciplines, providing CE to 200+ practicing professionals yearly, conducting diagnostic evaluations, and providing TA to Title V programs and community organizations. The program's 25-year history and track record (24,000+ trainees and 433,000 CE participants in FY23) demonstrate sustained impact. Dissemination through WVUCED network of 5,000+ contacts plus AUCD and WVU platforms ensures broad reach.\n\nResources and capabilities are strong with integration within WVUCED (federally designated UCEDD), qualified leadership (PD with EdD and 15+ years disability experience at 30% effort, Co-PD board-certified pediatric neurologist with 20 years LEND director experience), core faculty from required disciplines, appropriate staffing, and access to extensive clinical infrastructure. Institutional commitment is demonstrated through cost-sharing and letters of support.\n\nThe budget of $625,160 Year 1 and $626,920 Years 2-5 (total $3,132,840) is reasonable and well-justified, within Tier 3 ceiling, with appropriate allocation to personnel, trainee support ($198,000 Year 1), consultant services, and other costs. Indirect costs correctly calculated at 8% of MTDC.\n\nMinor weaknesses include limited detail on specific screening/diagnostic tools to be taught and competency assessment methods, incomplete medium-term trainee selection criteria, lack of specific strategies for recruiting racial/ethnic minority trainees, limited quantitative benchmarks for some evaluation objectives, PD effort at minimum 30% level, and lack of specific projections for number of diagnostic evaluations. These do not substantially diminish the overall strength of the application.\n\nThe application demonstrates excellent alignment with HRSA priorities for rural workforce development, interdisciplinary training, lived experience inclusion, and systems improvement for children/youth with autism/DD. Recommend for funding as requested.",
    "final_score": 93,
    "maximum_score": 100,
    "review_status": "ai_draft_human_validation_required",
    "certification": "Claude-generated draft. A human reviewer must verify every finding, citation, score, and budget recommendation."
}

# ── Helper functions ─────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_bold_para(doc, label, value, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run_label = p.add_run(label + ": ")
    run_label.bold = True
    p.add_run(value)
    return p


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3 + 0.25 * indent_level)
    p.add_run(text)
    return p


def pages_str(pages):
    if not pages:
        return ""
    return "  [p. " + ", ".join(str(pg) for pg in pages) + "]"


# ── Build document ───────────────────────────────────────────────────────────

doc = Document()

# Set narrow margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# ── Title block ──────────────────────────────────────────────────────────────

title = doc.add_heading(DATA["applicant_name"].upper(), 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(16)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("HRSA Grant Application — Peer Review Results").bold = True

doc.add_paragraph()

# ── Application summary box ──────────────────────────────────────────────────

add_bold_para(doc, "Applicant", DATA["applicant_name"])
add_bold_para(doc, "Application Number", DATA["application_number"])
add_bold_para(doc, "Agency", DATA["agency"])

total_score = DATA["final_score"]
max_score = DATA["maximum_score"]
score_p = doc.add_paragraph()
score_run = score_p.add_run(f"TOTAL SCORE: {total_score} / {max_score}")
score_run.bold = True
score_run.font.size = Pt(13)
score_run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9C)

doc.add_paragraph()

# ── Score summary table ──────────────────────────────────────────────────────

add_heading(doc, "SCORE SUMMARY", level=1)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Criterion"
hdr[1].text = "Score"
hdr[2].text = "Maximum"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

for c in DATA["criteria"]:
    row = table.add_row().cells
    row[0].text = c["name"]
    row[1].text = str(c["score"])
    row[2].text = str(c["maximum_points"])

# Totals row
tot = table.add_row().cells
tot[0].paragraphs[0].add_run("TOTAL").bold = True
tot[1].paragraphs[0].add_run(str(total_score)).bold = True
tot[2].paragraphs[0].add_run(str(max_score)).bold = True

doc.add_paragraph()

# ── Budget recommendation ────────────────────────────────────────────────────

add_heading(doc, "BUDGET RECOMMENDATION", level=1)

budget = DATA["budget"]
rec = "AS REQUESTED" if budget["recommendation"] == "as_requested" else "REDUCED"
p = doc.add_paragraph()
p.add_run(f"Recommendation: {rec}").bold = True

funds = budget["annual_recommended_funding"]
for i, amt in enumerate(funds, 1):
    doc.add_paragraph(f"Year {i}: ${amt:,.0f}", style="List Bullet")

doc.add_paragraph(budget["reduction_rationale"])

doc.add_paragraph()

# ── Overview sections ────────────────────────────────────────────────────────

add_heading(doc, "APPLICATION OVERVIEW", level=1)

ov = DATA["overview"]
sections = [
    ("Applicant Information", ov["applicant_information"]),
    ("Target Population", ov["target_population"]),
    ("Project Description", ov["project_description"]),
    ("Goals and Objectives", ov["goals_objectives"]),
    ("Significant Findings", ov["significant_findings"]),
    ("Other Information", ov["other_information"]),
]
for label, text in sections:
    add_heading(doc, label, level=2)
    doc.add_paragraph(text)

doc.add_paragraph()

# ── Detailed criteria ────────────────────────────────────────────────────────

add_heading(doc, "DETAILED CRITERIA REVIEW", level=1)

for c in DATA["criteria"]:
    # Criterion header
    crit_head = doc.add_heading(
        f"{c['name']}  —  Score: {c['score']} / {c['maximum_points']}", level=2
    )

    # Score rationale
    add_heading(doc, "Score Rationale", level=3)
    doc.add_paragraph(c["score_rationale"])

    # Strengths
    if c.get("strengths"):
        add_heading(doc, "Strengths", level=3)
        for s in c["strengths"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run(s["comment"] + pages_str(s.get("pages", [])))

    # Meets requirements
    if c.get("mets"):
        add_heading(doc, "Meets Requirements", level=3)
        for m in c["mets"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run(m["comment"] + pages_str(m.get("pages", [])))

    # Weaknesses
    if c.get("weaknesses"):
        add_heading(doc, "Weaknesses", level=3)
        for w in c["weaknesses"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run(w["comment"] + pages_str(w.get("pages", [])))

    doc.add_paragraph()

# ── Overall summary ──────────────────────────────────────────────────────────

add_heading(doc, "OVERALL SUMMARY", level=1)
doc.add_paragraph(DATA["overall_summary"])

doc.add_paragraph()

# ── Certification ────────────────────────────────────────────────────────────

add_heading(doc, "CERTIFICATION", level=1)
cert_p = doc.add_paragraph(DATA["certification"])
for run in cert_p.runs:
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True

# ── Save ─────────────────────────────────────────────────────────────────────

out_path = r"C:\Users\adeto\Downloads\HRSA 026-19\WEST VIRGINIA UNIVERSITY RESEARCH CORPORATION - Review Results.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
